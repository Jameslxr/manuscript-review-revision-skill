#!/usr/bin/env python3
"""Enforce continuous Word line numbering and dynamic page numbers in DOCX."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as exc:  # pragma: no cover - environment-dependent error path
    print(
        "python-docx is required. Use the bundled workspace Python runtime.",
        file=sys.stderr,
    )
    raise SystemExit(2) from exc


PAGE_FIELD_RE = re.compile(r"(?:^|\s)PAGE(?:\s|\\|$)", re.IGNORECASE)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Add continuous line numbering to every DOCX section and a dynamic "
            "PAGE field to every active header/footer story."
        )
    )
    parser.add_argument("document", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument(
        "--page-number-position",
        choices=("upper-right", "lower-center"),
        default="upper-right",
        help=(
            "Place dynamic PAGE fields in the upper-right header for the "
            "journal-neutral profile (default), or in the lower-center footer "
            "when an exact journal/template requires it."
        ),
    )
    return parser.parse_args()


def field_instruction(node: Any) -> str:
    if node.tag == qn("w:fldSimple"):
        return str(node.get(qn("w:instr")) or "")
    return str(node.text or "")


def story_has_page_field(story: Any) -> bool:
    for node in story._element.xpath(".//w:instrText | .//w:fldSimple"):
        if PAGE_FIELD_RE.search(field_instruction(node)):
            return True
    return False


def paragraph_has_page_field(paragraph: Any) -> bool:
    for node in paragraph._p.xpath(".//w:instrText | .//w:fldSimple"):
        if PAGE_FIELD_RE.search(field_instruction(node)):
            return True
    return False


def paragraph_has_only_page_field(paragraph: Any) -> bool:
    if not paragraph_has_page_field(paragraph):
        return False
    if paragraph._p.xpath(".//w:drawing | .//w:object | .//w:pict"):
        return False
    visible = paragraph.text.strip()
    return not visible or bool(re.fullmatch(r"\d+", visible))


def append_page_field(paragraph: Any, alignment: Any) -> None:
    paragraph.alignment = alignment
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    display = OxmlElement("w:t")
    display.text = "1"
    run._r.append(display)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def align_existing_page_fields(story: Any, alignment: Any) -> bool:
    found = False
    for paragraph in story.paragraphs:
        if paragraph_has_page_field(paragraph):
            paragraph.alignment = alignment
            found = True
    return found


def remove_field_only_page_paragraphs(story: Any) -> None:
    for paragraph in list(story.paragraphs):
        if not paragraph_has_only_page_field(paragraph):
            continue
        parent = paragraph._p.getparent()
        if parent is not None:
            parent.remove(paragraph._p)


def ensure_page_field(header: Any, footer: Any, position: str) -> None:
    if position == "upper-right":
        target, other = header, footer
        alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        target, other = footer, header
        alignment = WD_ALIGN_PARAGRAPH.CENTER

    if align_existing_page_fields(target, alignment):
        return

    remove_field_only_page_paragraphs(other)
    if story_has_page_field(other):
        raise ValueError(
            "Unable to relocate a PAGE field embedded with other header/footer "
            "content. Move it manually or supply a clean journal template."
        )

    paragraph = target.paragraphs[-1]
    if paragraph.text.strip() or paragraph._p.xpath(
        ".//w:drawing | .//w:object | .//w:pict | .//w:fldChar | .//w:instrText"
    ):
        paragraph = target.add_paragraph()
    append_page_field(paragraph, alignment)


def enforce_line_numbering(document: Any) -> None:
    for section in document.sections:
        section_properties = section._sectPr
        line_number_nodes = section_properties.xpath("./w:lnNumType")
        line_numbers = (
            line_number_nodes[0]
            if line_number_nodes
            else OxmlElement("w:lnNumType")
        )
        if not line_number_nodes:
            section_properties.append(line_numbers)
        for duplicate in line_number_nodes[1:]:
            section_properties.remove(duplicate)
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:restart"), "continuous")
        line_numbers.set(qn("w:distance"), "360")
        line_numbers.attrib.pop(qn("w:start"), None)

    for suppress_node in list(
        document.element.body.xpath(".//w:pPr/w:suppressLineNumbers")
    ):
        parent = suppress_node.getparent()
        if parent is not None:
            parent.remove(suppress_node)


def enforce_page_numbering(document: Any, position: str) -> None:
    even_and_odd = (
        document.settings.element.find(qn("w:evenAndOddHeaders")) is not None
    )
    for section in document.sections:
        page_number_nodes = section._sectPr.xpath("./w:pgNumType")
        for page_number_node in page_number_nodes:
            page_number_node.attrib.pop(qn("w:start"), None)

        ensure_page_field(section.header, section.footer, position)
        if section.different_first_page_header_footer:
            ensure_page_field(
                section.first_page_header,
                section.first_page_footer,
                position,
            )
        if even_and_odd:
            ensure_page_field(
                section.even_page_header,
                section.even_page_footer,
                position,
            )


def enforce(document_path: Path, output_path: Path, page_number_position: str) -> None:
    if document_path.resolve() == output_path.resolve():
        raise ValueError("--out must differ from the input path; preserve the source DOCX.")
    document = Document(str(document_path))
    enforce_line_numbering(document)
    enforce_page_numbering(document, page_number_position)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    args = parse_args()
    try:
        enforce(args.document, args.out, args.page_number_position)
    except (OSError, ValueError) as exc:
        print(f"Unable to enforce DOCX numbering: {exc}", file=sys.stderr)
        return 2
    print(f"DOCX numbering enforced: {args.out.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
