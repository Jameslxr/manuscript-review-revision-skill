#!/usr/bin/env python3
"""Regression tests for manuscript-wide semantic vertical rhythm."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"


def run_script(name: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        text=True,
        capture_output=True,
        check=False,
    )


def add_style(document: Document, name: str, size: float) -> object:
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.paragraph_format.space_after = Pt(9)
    style.paragraph_format.line_spacing = 1.15
    return style


def build_semantic_stress_sample(path: Path) -> list[str]:
    document = Document()
    styles = {
        "title": add_style(document, "Manuscript Title", 18),
        "authors": add_style(document, "Manuscript Authors", 8),
        "affiliation": add_style(document, "Manuscript Affiliation", 8),
        "author_note": add_style(document, "Manuscript Author Note", 8),
        "correspondence": add_style(document, "Manuscript Correspondence", 9),
        "orcid": add_style(document, "Manuscript ORCID", 9),
    }
    entries = (
        ("Metastasis in Hepatocellular Carcinoma", styles["title"]),
        ("James Li1 and Alex Smith1,*", styles["authors"]),
        ("1 Department of Pathology, Example University", styles["affiliation"]),
        ("These authors contributed equally.", styles["author_note"]),
        ("*Correspondence: alex.smith@example.edu", styles["correspondence"]),
        ("ORCID: https://orcid.org/0000-0001-2345-6789", styles["orcid"]),
    )
    for text, style in entries:
        paragraph = document.add_paragraph(text, style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_heading("Abstract", level=1)
    document.add_paragraph("Abstract body without standardized spacing.")
    document.add_paragraph("Keywords: liver cancer; metastasis")
    document.add_heading("Introduction", level=1)
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("First main-text paragraph.")
    document.add_heading("2. Metastatic programs", level=1)
    document.add_paragraph("Second-section body paragraph.")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_heading("2.1 Vascular invasion", level=2)
    document.add_paragraph("Subsection body paragraph.")
    document.add_paragraph("CRediT authorship contribution statement")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph(
        "J.L.: Conceptualization, Methodology, Writing – original draft."
    )
    document.add_paragraph("")
    document.add_paragraph(
        "A.S.: Supervision, Writing – review & editing."
    )
    document.add_paragraph("Funding: No external funding was received.")
    document.add_heading("References", level=1)
    document.add_paragraph("1. Reference list retained as supplied.")
    document.add_paragraph("2. A second reference remains consecutive.")
    original = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    document.save(path)
    return original


def apply_profile(
    source: Path, output: Path, spacing: str = "double"
) -> subprocess.CompletedProcess[str]:
    return run_script(
        "apply_manuscript_profile.py",
        source,
        "--out",
        output,
        "--line-spacing",
        spacing,
        "--body-style",
        "Normal",
        "--title-style",
        "Manuscript Title",
        "--authors-style",
        "Manuscript Authors",
        "--affiliation-style",
        "Manuscript Affiliation",
        "--author-note-style",
        "Manuscript Author Note",
        "--correspondence-style",
        "Manuscript Correspondence",
        "--orcid-style",
        "Manuscript ORCID",
    )


def audit_semantic(path: Path, report: Path) -> dict[str, object]:
    run_script(
        "audit_docx_semantic_rhythm.py",
        path,
        "--expected-line-spacing",
        "double",
        "--output-json",
        report,
    )
    return json.loads(report.read_text(encoding="utf-8"))


class SemanticRhythmTests(unittest.TestCase):
    def test_one_point_five_resolves_across_every_manuscript_role(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "normalized.docx"
            build_semantic_stress_sample(source)
            result = apply_profile(source, output, "1.5")
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            audited = run_script(
                "audit_docx_semantic_rhythm.py",
                output,
                "--expected-line-spacing",
                "1.5",
                "--json",
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            self.assertEqual(
                json.loads(audited.stdout)["status"], "SEMANTIC_RHYTHM_PASS"
            )

    def test_profile_repairs_global_typography_and_semantic_boundaries(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "normalized.docx"
            release = root / "release.docx"
            original = build_semantic_stress_sample(source)
            result = apply_profile(source, output)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

            report = audit_semantic(output, root / "semantic.json")
            self.assertEqual(report["status"], "SEMANTIC_RHYTHM_PASS")
            numbered = run_script(
                "enforce_docx_line_page_numbers.py", output, "--out", release
            )
            self.assertEqual(numbered.returncode, 0, numbered.stdout + numbered.stderr)
            structural_path = root / "structural.json"
            structural = run_script(
                "audit_docx_manuscript_style.py",
                release,
                "--expected-line-spacing",
                "double",
                "--body-style",
                "Manuscript Body",
                "--output-json",
                structural_path,
            )
            self.assertEqual(structural.returncode, 0, structural.stdout + structural.stderr)
            structural_report = json.loads(structural_path.read_text(encoding="utf-8"))
            self.assertEqual(structural_report["inspected_title_heading_count"], 7)
            document = Document(output)
            self.assertEqual(
                original,
                [paragraph.text for paragraph in document.paragraphs if paragraph.text],
            )

            by_text = {paragraph.text: paragraph for paragraph in document.paragraphs}
            for text in (
                "James Li1 and Alex Smith1,*",
                "1 Department of Pathology, Example University",
                "*Correspondence: alex.smith@example.edu",
                "Keywords: liver cancer; metastasis",
                "2. Metastatic programs",
                "2.1 Vascular invasion",
                "CRediT authorship contribution statement",
            ):
                paragraph = by_text[text]
                self.assertAlmostEqual(float(paragraph.paragraph_format.line_spacing), 2.0)
                for run in paragraph.runs:
                    if run.text.strip():
                        self.assertAlmostEqual(run.font.size.pt, 12.0)

            keywords = by_text["Keywords: liver cancer; metastasis"]
            self.assertTrue(keywords.runs[0].bold)
            self.assertFalse(any(run.bold for run in keywords.runs[1:] if run.text.strip()))
            first_reference = by_text["1. Reference list retained as supplied."]
            second_reference = by_text["2. A second reference remains consecutive."]
            self.assertEqual(first_reference.style.name, "Manuscript Reference")
            self.assertEqual(second_reference.style.name, "Manuscript Reference")
            first_index = next(
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph.text == first_reference.text
            )
            self.assertEqual(
                document.paragraphs[first_index + 1].text, second_reference.text
            )
            first_credit = by_text[
                "J.L.: Conceptualization, Methodology, Writing – original draft."
            ]
            second_credit = by_text[
                "A.S.: Supervision, Writing – review & editing."
            ]
            first_credit_index = next(
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph.text == first_credit.text
            )
            self.assertEqual(
                document.paragraphs[first_credit_index + 1].text,
                second_credit.text,
            )
            self.assertEqual(first_credit.style.name, "Manuscript CRediT Entry")
            self.assertEqual(second_credit.style.name, "Manuscript CRediT Entry")

    def test_custom_ten_point_body_and_table_are_normalized_to_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "custom-source.docx"
            output = root / "custom-normalized.docx"
            document = Document()
            custom = add_style(document, "Journal Main Text", 10)
            title = document.add_paragraph("Custom-style manuscript")
            title.style = document.styles["Title"]
            title.runs[0].font.size = Pt(18)
            document.add_heading("Abstract", level=1)
            document.add_paragraph("Custom abstract body.", style=custom)
            document.add_heading("Introduction", level=1)
            document.add_paragraph("Custom main-text body.", style=custom)
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).paragraphs[0].add_run("Custom table text.")
            table.cell(0, 0).paragraphs[0].style = custom
            document.save(source)

            result = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                output,
                "--title-paragraph",
                "1",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            normalized = Document(output)
            for paragraph in normalized.paragraphs[1:]:
                for run in paragraph.runs:
                    if run.text.strip():
                        self.assertAlmostEqual(run.font.size.pt, 12.0)
            table_run = normalized.tables[0].cell(0, 0).paragraphs[0].runs[0]
            self.assertAlmostEqual(table_run.font.size.pt, 12.0)

            audited = run_script(
                "audit_docx_semantic_rhythm.py",
                output,
                "--expected-line-spacing",
                "double",
                "--json",
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)

    def test_explicit_journal_typography_values_are_parameterized(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "normalized.docx"
            build_semantic_stress_sample(source)
            result = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                output,
                "--line-spacing",
                "1.5",
                "--font-name",
                "Times New Roman",
                "--body-font-size",
                "10",
                "--title-font-size",
                "15",
                "--table-font-size",
                "10",
                "--body-style",
                "Normal",
                "--title-style",
                "Manuscript Title",
                "--authors-style",
                "Manuscript Authors",
                "--affiliation-style",
                "Manuscript Affiliation",
                "--author-note-style",
                "Manuscript Author Note",
                "--correspondence-style",
                "Manuscript Correspondence",
                "--orcid-style",
                "Manuscript ORCID",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            audited = run_script(
                "audit_docx_semantic_rhythm.py",
                output,
                "--expected-line-spacing",
                "1.5",
                "--expected-body-font-size",
                "10",
                "--expected-title-font-size",
                "15",
                "--expected-table-font-size",
                "10",
                "--json",
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)

    def test_semantic_audit_cannot_skip_unclassified_ten_point_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            broken = root / "broken.docx"
            build_semantic_stress_sample(source)
            result = apply_profile(source, normalized)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            document = Document(normalized)
            target = next(
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text == "First main-text paragraph."
            )
            custom = add_style(document, "Journal Main Text", 10)
            custom.font.name = "Times New Roman"
            custom.paragraph_format.space_before = Pt(0)
            custom.paragraph_format.space_after = Pt(0)
            custom.paragraph_format.line_spacing = 2.0
            target.style = custom
            for run in target.runs:
                run.font.size = None
                run.font.name = None
            document.save(broken)
            audited = run_script(
                "audit_docx_semantic_rhythm.py",
                broken,
                "--expected-line-spacing",
                "double",
                "--json",
            )
            self.assertEqual(audited.returncode, 1)
            report = json.loads(audited.stdout)
            self.assertIn(
                "VISIBLE_FONT_SIZE_MISMATCH",
                {issue["code"] for issue in report["issues"]},
            )

    def test_semantic_audit_rejects_each_reported_regression_class(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            broken = root / "broken.docx"
            build_semantic_stress_sample(source)
            result = apply_profile(source, normalized)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

            document = Document(normalized)
            by_text = {paragraph.text: paragraph for paragraph in document.paragraphs}
            affiliation = by_text["1 Department of Pathology, Example University"]
            affiliation.runs[0].font.size = Pt(9)
            affiliation.paragraph_format.line_spacing = 1.0
            keywords = by_text["Keywords: liver cancer; metastasis"]
            keywords.runs[0].font.bold = False
            section = by_text["2. Metastatic programs"]
            previous = section._p.getprevious()
            if previous is not None:
                previous.getparent().remove(previous)
            section._p.addnext(OxmlElement("w:p"))
            first_credit = by_text[
                "J.L.: Conceptualization, Methodology, Writing – original draft."
            ]
            second_credit = by_text[
                "A.S.: Supervision, Writing – review & editing."
            ]
            first_credit._p.addnext(OxmlElement("w:p"))
            second_credit.runs[0].text = "A.S.: Oversight."
            document.save(broken)

            report = audit_semantic(broken, root / "broken.json")
            self.assertEqual(report["status"], "FAIL")
            codes = {issue["code"] for issue in report["issues"]}
            self.assertIn("BODY_SIZE_ROLE_MISMATCH", codes)
            self.assertIn("GLOBAL_LINE_SPACING_MISMATCH", codes)
            self.assertIn("KEYWORDS_LABEL_NOT_BOLD", codes)
            self.assertIn("SEMANTIC_BLANK_BEFORE", codes)
            self.assertIn("HEADING_BLANK_AFTER", codes)
            self.assertIn("CREDIT_ENTRY_BLANK_BETWEEN", codes)
            self.assertIn("CREDIT_ENTRY_WITHOUT_STANDARD_ROLE", codes)


if __name__ == "__main__":
    unittest.main()
