#!/usr/bin/env python3
"""Focused regression tests for the standalone DOCX formatting skill."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"


def run_script(name: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        text=True,
        capture_output=True,
        check=False,
    )


def add_dynamic_page_field(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, display, end):
        run._r.append(node)


def apply_required_numbering(document: Document) -> None:
    for section in document.sections:
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:restart"), "continuous")
        section._sectPr.append(line_numbers)
        add_dynamic_page_field(section.footer.paragraphs[0])


def set_double_zero_spacing(document: Document) -> None:
    normal = document.styles["Normal"].paragraph_format
    normal.space_before = Pt(0)
    normal.space_after = Pt(0)
    normal.line_spacing = 2.0


class StandaloneFormattingTests(unittest.TestCase):
    def test_missing_literal_blank_and_nonzero_spacing_fail(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "bad-spacing.docx"
            document = Document()
            normal = document.styles["Normal"].paragraph_format
            normal.space_before = Pt(0)
            normal.space_after = Pt(8)
            normal.line_spacing = 2.0
            document.add_paragraph("First paragraph.")
            document.add_paragraph("Second paragraph.")
            apply_required_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BODY_PARAGRAPH_SPACE_AFTER", result.stdout)
            self.assertIn("MISSING_LITERAL_BLANK_PARAGRAPH", result.stdout)

    def test_literal_blank_zero_spacing_and_numbering_pass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "numbered.docx"
            report = root / "audit.json"
            document = Document()
            set_double_zero_spacing(document)
            document.add_paragraph("First paragraph.")
            separator = document.add_paragraph("")
            separator.paragraph_format.space_before = Pt(0)
            separator.paragraph_format.space_after = Pt(0)
            separator.paragraph_format.line_spacing = 2.0
            document.add_paragraph("Second paragraph.")
            document.save(source)

            enforced = run_script(
                "enforce_docx_line_page_numbers.py", source, "--out", output
            )
            self.assertEqual(
                enforced.returncode, 0, enforced.stdout + enforced.stderr
            )
            audited = run_script(
                "audit_docx_manuscript_style.py",
                output,
                "--expected-line-spacing",
                "double",
                "--output-json",
                report,
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            self.assertIn("MECHANICAL_PASS", audited.stdout)
            self.assertTrue(report.exists())

    def test_blank_separator_with_spaces_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "blank-with-spaces.docx"
            document = Document()
            set_double_zero_spacing(document)
            document.add_paragraph("First paragraph.")
            separator = document.add_paragraph("   ")
            separator.paragraph_format.space_before = Pt(0)
            separator.paragraph_format.space_after = Pt(0)
            separator.paragraph_format.line_spacing = 2.0
            document.add_paragraph("Second paragraph.")
            apply_required_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BLANK_PARAGRAPH_CONTAINS_WHITESPACE", result.stdout)

    def test_missing_line_and_page_numbering_fail(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "unnumbered.docx"
            document = Document()
            set_double_zero_spacing(document)
            document.add_paragraph("Body text.")
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("SECTION_LINE_NUMBERING_MISSING", result.stdout)
            self.assertIn("PAGE_NUMBER_FIELD_MISSING", result.stdout)

    def test_numbering_enforcer_covers_active_multisection_stories(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "numbered.docx"
            document = Document()
            set_double_zero_spacing(document)
            document.sections[0].different_first_page_header_footer = True
            document.add_paragraph("First section body.")
            document.add_section(WD_SECTION.NEW_PAGE)
            document.sections[1].different_first_page_header_footer = True
            document.add_paragraph("Second section body.")
            document.settings.element.append(OxmlElement("w:evenAndOddHeaders"))
            document.save(source)

            enforced = run_script(
                "enforce_docx_line_page_numbers.py", source, "--out", output
            )
            self.assertEqual(enforced.returncode, 0)
            audited = run_script("audit_docx_manuscript_style.py", output)
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            self.assertIn("Continuous line-numbered sections: 2/2", audited.stdout)
            self.assertIn(
                "Dynamic PAGE fields in active page stories: 6/6", audited.stdout
            )

    def test_hidden_word_autospacing_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "hidden-autospacing.docx"
            document = Document()
            set_double_zero_spacing(document)
            style_properties = document.styles["Normal"].element.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:afterAutospacing"), "1")
            style_properties.append(spacing)
            document.add_paragraph("First paragraph.")
            document.add_paragraph("")
            document.add_paragraph("Second paragraph.")
            apply_required_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BODY_PARAGRAPH_AUTOSPACING_AFTER", result.stdout)

    def test_custom_style_requires_explicit_classification(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "custom-style.docx"
            document = Document()
            set_double_zero_spacing(document)
            custom = document.styles.add_style(
                "Custom Front Matter", WD_STYLE_TYPE.PARAGRAPH
            )
            custom.base_style = None
            document.add_paragraph("Custom title", style=custom)
            document.add_paragraph("Body text.")
            apply_required_numbering(document)
            document.save(path)

            blocked = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(blocked.returncode, 1)
            self.assertIn("UNCLASSIFIED_NONEMPTY_PARAGRAPH_STYLE", blocked.stdout)

            classified = run_script(
                "audit_docx_manuscript_style.py",
                path,
                "--exclude-style",
                "Custom Front Matter",
            )
            self.assertEqual(
                classified.returncode,
                0,
                classified.stdout + classified.stderr,
            )


if __name__ == "__main__":
    unittest.main()
