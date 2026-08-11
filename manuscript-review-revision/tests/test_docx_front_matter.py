#!/usr/bin/env python3
"""Benchmark the journal-neutral front-matter profile and repair path."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
ROLE_ARGS = (
    "--title-style",
    "Manuscript Title",
    "--authors-style",
    "Manuscript Authors",
    "--affiliation-style",
    "Manuscript Affiliation",
    "--author_note-style",
    "Manuscript Author Note",
    "--correspondence-style",
    "Manuscript Correspondence",
    "--orcid-style",
    "Manuscript ORCID",
)


def run_script(name: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        text=True,
        capture_output=True,
        check=False,
    )


def add_style(
    document: Document, name: str, size: float, *, bold: bool = False
) -> object:
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = None
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 2.0
    return style


def build_sample(
    path: Path,
    *,
    alignments: tuple[object, object, object, object] | None = None,
    title_size: float = 15,
    title_author_blanks: int = 1,
    author_affiliation_blanks: int = 1,
    affiliation_correspondence_blanks: int = 1,
    extra_blanks: int = 1,
    table_front_matter: bool = False,
) -> None:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = {
        "title": add_style(document, "Manuscript Title", title_size, bold=True),
        "authors": add_style(document, "Manuscript Authors", 12),
        "affiliation": add_style(document, "Manuscript Affiliation", 12),
        "correspondence": add_style(
            document, "Manuscript Correspondence", 12
        ),
        "keywords": add_style(document, "Manuscript Keywords", 12),
        "body": add_style(document, "Manuscript Body", 12),
    }
    styles["body"].paragraph_format.line_spacing = 2.0
    entries = (
        ("Natural Manuscript Formatting Benchmark", styles["title"]),
        ("James Li1 and Alex Smith1,*", styles["authors"]),
        (
            "1 Department of Pathology, Example University, Chicago, IL, USA",
            styles["affiliation"],
        ),
        ("*Correspondence: alex.smith@example.edu", styles["correspondence"]),
    )

    if table_front_matter:
        table = document.add_table(rows=4, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row, (text, style) in zip(table.rows, entries):
            paragraph = row.cells[0].paragraphs[0]
            paragraph.style = style
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(text)
        document.add_paragraph("", style=styles["body"])
    else:
        assert alignments is not None
        gap_counts = (
            title_author_blanks,
            author_affiliation_blanks,
            affiliation_correspondence_blanks,
            0,
        )
        for (text, style), alignment, gap_count in zip(
            entries, alignments, gap_counts
        ):
            paragraph = document.add_paragraph(text, style=style)
            paragraph.alignment = alignment
            for _ in range(gap_count):
                document.add_paragraph("", style=styles["body"])
        for _ in range(extra_blanks):
            document.add_paragraph("", style=styles["body"])

    heading = document.add_heading("Abstract", level=1)
    heading.style.font.name = "Times New Roman"
    heading.style.font.size = Pt(12)
    heading.style.font.bold = True
    heading.style.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.line_spacing = 2.0
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph("First body paragraph.", style=styles["body"])
    separator = document.add_paragraph("", style=styles["body"])
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(0)
    separator.paragraph_format.line_spacing = 2.0
    document.add_paragraph("Second body paragraph.", style=styles["body"])
    document.add_paragraph(
        "Keywords: liver cancer; metastasis", style=styles["keywords"]
    )
    document.save(path)


def build_full_semantic_matrix_sample(path: Path) -> None:
    document = Document()
    styles = {
        "title": add_style(document, "Manuscript Title", 15, bold=True),
        "authors": add_style(document, "Manuscript Authors", 12),
        "affiliation": add_style(document, "Manuscript Affiliation", 12),
        "author_note": add_style(document, "Manuscript Author Note", 12),
        "correspondence": add_style(
            document, "Manuscript Correspondence", 12
        ),
        "orcid": add_style(document, "Manuscript ORCID", 12),
        "body": add_style(document, "Manuscript Body", 12),
    }
    sequence = (
        ("Natural Manuscript Formatting Benchmark", "title"),
        ("James Li1 and Alex Smith1,*", "authors"),
        ("", "body"),
        ("", "body"),
        ("1 Department of Pathology, Example University", "affiliation"),
        ("", "body"),
        ("2 Department of Medicine, Example University", "affiliation"),
        ("These authors contributed equally and share first authorship.", "author_note"),
        ("*Correspondence: alex.smith@example.edu", "correspondence"),
        ("", "body"),
        ("ORCID: James Li, https://orcid.org/0000-0001-2345-6789", "orcid"),
        ("", "body"),
        ("", "body"),
    )
    for text, role in sequence:
        document.add_paragraph(text, style=styles[role])
    heading = document.add_paragraph("Abstract", style=styles["body"])
    heading.style = document.styles["Manuscript Body"]
    document.add_paragraph("First body paragraph.", style=styles["body"])
    document.save(path)


def front_matter_text_sequence(path: Path) -> list[str]:
    values: list[str] = []
    for paragraph in Document(path).paragraphs:
        values.append(paragraph.text)
        if paragraph.text.strip().casefold() == "abstract":
            break
    return values


def enforce_numbering(source: Path, output: Path, *args: object) -> None:
    result = run_script(
        "enforce_docx_line_page_numbers.py", source, "--out", output, *args
    )
    if result.returncode:
        raise AssertionError(result.stdout + result.stderr)


def build_all_normal_sample(path: Path) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 78, 121)
    normal.paragraph_format.space_after = Pt(8)
    for text in (
        "Natural Manuscript Formatting Benchmark",
        "",
        "James Li1 and Alex Smith1,*",
        "1 Department of Pathology, Example University, Chicago, IL, USA",
        "*Correspondence: alex.smith@example.edu",
        "",
        "Abstract",
        "First body paragraph.",
        "Second body paragraph.",
    ):
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)


def front_audit(path: Path, output_json: Path, *args: object) -> dict[str, object]:
    run_script(
        "audit_docx_front_matter.py",
        path,
        *ROLE_ARGS,
        *args,
        "--output-json",
        output_json,
    )
    return json.loads(output_json.read_text(encoding="utf-8"))


class FrontMatterBenchmarkTests(unittest.TestCase):
    def test_all_normal_draft_can_be_repaired_by_paragraph_roles(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "all-normal.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            build_all_normal_sample(source)
            role_paragraphs = (
                "--title-paragraph",
                "1",
                "--authors-paragraph",
                "3",
                "--affiliation-paragraph",
                "4",
                "--correspondence-paragraph",
                "5",
            )
            applied = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                normalized,
                "--body-style",
                "Normal",
                *role_paragraphs,
            )
            self.assertEqual(applied.returncode, 0, applied.stdout + applied.stderr)
            enforce_numbering(normalized, release)
            audited = front_audit(release, root / "front.json")
            self.assertEqual(audited["status"], "FRONT_MATTER_PASS")

    def test_gold_left_profile_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "numbered.docx"
            report = root / "front.json"
            left = WD_ALIGN_PARAGRAPH.LEFT
            build_sample(source, alignments=(left, left, left, left))
            enforce_numbering(source, output)
            audited = front_audit(output, report)
            self.assertEqual(audited["status"], "FRONT_MATTER_PASS")

    def test_seven_adversarial_front_matters_are_rejected(self) -> None:
        left = WD_ALIGN_PARAGRAPH.LEFT
        center = WD_ALIGN_PARAGRAPH.CENTER
        right = WD_ALIGN_PARAGRAPH.RIGHT
        cases = {
            "centered": {"alignments": (center, center, center, center)},
            "mixed": {"alignments": (left, center, right, center)},
            "oversized": {
                "alignments": (left, left, left, left),
                "title_size": 26,
            },
            "missing-title-author-blank": {
                "alignments": (left, left, left, left),
                "title_author_blanks": 0,
            },
            "duplicate-title-author-blank": {
                "alignments": (left, left, left, left),
                "title_author_blanks": 2,
            },
            "extra-blanks": {
                "alignments": (left, left, left, left),
                "extra_blanks": 4,
            },
            "table": {"table_front_matter": True},
        }
        expected_codes = {
            "centered": "TITLE_ALIGNMENT",
            "mixed": "AUTHORS_ALIGNMENT",
            "oversized": "TITLE_FONT_SIZE",
            "missing-title-author-blank": "FRONT_MATTER_BLOCK_BLANK_COUNT",
            "duplicate-title-author-blank": "FRONT_MATTER_BLOCK_BLANK_COUNT",
            "extra-blanks": "FRONT_MATTER_BLOCK_BLANK_COUNT",
            "table": "FRONT_MATTER_TABLE_LAYOUT",
        }
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            for name, kwargs in cases.items():
                with self.subTest(name=name):
                    source = root / f"{name}-source.docx"
                    output = root / f"{name}-numbered.docx"
                    report = root / f"{name}.json"
                    build_sample(source, **kwargs)
                    enforce_numbering(source, output)
                    audited = front_audit(output, report)
                    self.assertEqual(audited["status"], "FAIL")
                    codes = {issue["code"] for issue in audited["issues"]}
                    self.assertIn(expected_codes[name], codes)

    def test_profile_normalizer_repairs_all_seven_adversarial_cases(self) -> None:
        left = WD_ALIGN_PARAGRAPH.LEFT
        center = WD_ALIGN_PARAGRAPH.CENTER
        right = WD_ALIGN_PARAGRAPH.RIGHT
        cases = {
            "centered": {"alignments": (center, center, center, center)},
            "mixed": {"alignments": (left, center, right, center)},
            "oversized": {
                "alignments": (left, left, left, left),
                "title_size": 26,
            },
            "missing-title-author-blank": {
                "alignments": (left, left, left, left),
                "title_author_blanks": 0,
            },
            "duplicate-title-author-blank": {
                "alignments": (left, left, left, left),
                "title_author_blanks": 2,
            },
            "extra-blanks": {
                "alignments": (left, left, left, left),
                "extra_blanks": 4,
            },
            "table": {"table_front_matter": True},
        }
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            for name, kwargs in cases.items():
                with self.subTest(name=name):
                    source = root / f"{name}-source.docx"
                    normalized = root / f"{name}-normalized.docx"
                    release = root / f"{name}-release.docx"
                    build_sample(source, **kwargs)
                    applied = run_script(
                        "apply_manuscript_profile.py",
                        source,
                        "--out",
                        normalized,
                        "--body-style",
                        "Manuscript Body",
                        *ROLE_ARGS,
                    )
                    self.assertEqual(
                        applied.returncode, 0, applied.stdout + applied.stderr
                    )
                    enforce_numbering(normalized, release)
                    front = front_audit(release, root / f"{name}-front.json")
                    self.assertEqual(front["status"], "FRONT_MATTER_PASS")

                    style = run_script(
                        "audit_docx_manuscript_style.py",
                        release,
                        "--expected-line-spacing",
                        "double",
                        "--body-style",
                        "Manuscript Body",
                        "--exclude-style",
                        "Manuscript Title",
                        "--exclude-style",
                        "Manuscript Authors",
                        "--exclude-style",
                        "Manuscript Affiliation",
                        "--exclude-style",
                        "Manuscript Correspondence",
                        "--exclude-style",
                        "Manuscript Keywords",
                        "--exclude-style",
                        "Manuscript Heading",
                    )
                    self.assertEqual(style.returncode, 0, style.stdout + style.stderr)

    def test_legacy_compact_title_author_override_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            left = WD_ALIGN_PARAGRAPH.LEFT
            build_sample(source, alignments=(left, left, left, left))

            applied = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                root / "normalized.docx",
                "--title-author-gap",
                "compact",
                "--body-style",
                "Manuscript Body",
                *ROLE_ARGS,
            )
            self.assertEqual(applied.returncode, 2)
            self.assertIn("unrecognized arguments", applied.stderr)

    def test_full_semantic_matrix_repairs_every_present_block_boundary(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            build_full_semantic_matrix_sample(source)
            applied = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                normalized,
                "--body-style",
                "Manuscript Body",
                *ROLE_ARGS,
            )
            self.assertEqual(applied.returncode, 0, applied.stdout + applied.stderr)
            enforce_numbering(normalized, release)
            audited = front_audit(release, root / "front.json")
            self.assertEqual(audited["status"], "FRONT_MATTER_PASS")
            self.assertEqual(
                front_matter_text_sequence(release),
                [
                    "Natural Manuscript Formatting Benchmark",
                    "",
                    "James Li1 and Alex Smith1,*",
                    "",
                    "1 Department of Pathology, Example University",
                    "2 Department of Medicine, Example University",
                    "",
                    "These authors contributed equally and share first authorship.",
                    "",
                    "*Correspondence: alex.smith@example.edu",
                    "",
                    "ORCID: James Li, https://orcid.org/0000-0001-2345-6789",
                    "",
                    "Abstract",
                ],
            )
            self.assertEqual(audited["role_counts"]["author_note"], 1)
            self.assertEqual(audited["role_counts"]["orcid"], 1)

    def test_missing_optional_blocks_still_keep_one_natural_blank(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            release = root / "release.docx"
            left = WD_ALIGN_PARAGRAPH.LEFT
            build_sample(source, alignments=(left, left, left, left))
            enforce_numbering(source, release)
            audited = front_audit(release, root / "front.json")
            self.assertEqual(audited["status"], "FRONT_MATTER_PASS")
            self.assertEqual(audited["role_counts"]["author_note"], 0)
            self.assertEqual(audited["role_counts"]["orcid"], 0)

    def test_blank_inside_multi_paragraph_role_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            release = root / "release.docx"
            build_full_semantic_matrix_sample(source)
            enforce_numbering(source, release)
            audited = front_audit(release, root / "front.json")
            self.assertEqual(audited["status"], "FAIL")
            self.assertIn(
                "FRONT_MATTER_WITHIN_BLOCK_BLANK_COUNT",
                {issue["code"] for issue in audited["issues"]},
            )

    def test_page_number_position_profile_is_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            upper = root / "upper.docx"
            lower = root / "lower.docx"
            left = WD_ALIGN_PARAGRAPH.LEFT
            build_sample(source, alignments=(left, left, left, left))

            enforce_numbering(source, upper)
            upper_report = front_audit(upper, root / "upper.json")
            self.assertEqual(upper_report["status"], "FRONT_MATTER_PASS")

            enforce_numbering(
                source, lower, "--page-number-position", "lower-center"
            )
            lower_report = front_audit(
                lower,
                root / "lower.json",
                "--expected-page-number-position",
                "lower-center",
            )
            self.assertEqual(lower_report["status"], "FRONT_MATTER_PASS")

            wrong_profile = front_audit(lower, root / "wrong.json")
            self.assertEqual(wrong_profile["status"], "FAIL")
            codes = {issue["code"] for issue in wrong_profile["issues"]}
            self.assertIn("PAGE_NUMBER_POSITION", codes)

    def test_front_matter_blank_spacing_is_rejected_and_repaired(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            numbered = root / "numbered.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            left = WD_ALIGN_PARAGRAPH.LEFT
            build_sample(source, alignments=(left, left, left, left))

            document = Document(source)
            blank = next(
                paragraph for paragraph in document.paragraphs if not paragraph.text
            )
            blank.paragraph_format.space_after = Pt(8)
            document.save(source)
            enforce_numbering(source, numbered)
            failed = front_audit(numbered, root / "failed.json")
            self.assertEqual(failed["status"], "FAIL")
            self.assertIn(
                "FRONT_MATTER_BLANK_SPACE_AFTER",
                {issue["code"] for issue in failed["issues"]},
            )

            applied = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                normalized,
                "--body-style",
                "Manuscript Body",
                *ROLE_ARGS,
            )
            self.assertEqual(applied.returncode, 0, applied.stdout + applied.stderr)
            enforce_numbering(normalized, release)
            repaired = front_audit(release, root / "repaired.json")
            self.assertEqual(repaired["status"], "FRONT_MATTER_PASS")

    def test_normalizer_is_idempotent_and_preserves_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            once = root / "once.docx"
            twice = root / "twice.docx"
            center = WD_ALIGN_PARAGRAPH.CENTER
            build_sample(source, alignments=(center, center, center, center))
            source_text = [paragraph.text for paragraph in Document(source).paragraphs]

            first = run_script(
                "apply_manuscript_profile.py",
                source,
                "--out",
                once,
                "--body-style",
                "Manuscript Body",
                *ROLE_ARGS,
            )
            self.assertEqual(first.returncode, 0, first.stdout + first.stderr)
            second = run_script(
                "apply_manuscript_profile.py",
                once,
                "--out",
                twice,
                "--body-style",
                "Manuscript Body",
            )
            self.assertEqual(second.returncode, 0, second.stdout + second.stderr)

            self.assertEqual(
                [paragraph.text for paragraph in Document(once).paragraphs],
                [paragraph.text for paragraph in Document(twice).paragraphs],
            )
            self.assertEqual(
                [text for text in source_text if text],
                [
                    paragraph.text
                    for paragraph in Document(once).paragraphs
                    if paragraph.text
                ],
            )
            with ZipFile(once) as left_zip, ZipFile(twice) as right_zip:
                self.assertEqual(
                    left_zip.read("word/document.xml"),
                    right_zip.read("word/document.xml"),
                )
                self.assertEqual(
                    left_zip.read("word/styles.xml"),
                    right_zip.read("word/styles.xml"),
                )

    def test_release_validator_is_fail_closed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            structural = root / "structural.json"
            front = root / "front.json"
            semantic = root / "semantic.json"
            structural.write_text(
                json.dumps({"status": "MECHANICAL_PASS"}), encoding="utf-8"
            )
            front.write_text(
                json.dumps({"status": "FRONT_MATTER_PASS"}), encoding="utf-8"
            )
            semantic.write_text(
                json.dumps({"status": "SEMANTIC_RHYTHM_PASS"}), encoding="utf-8"
            )
            passed = run_script(
                "validate_format_release.py",
                structural,
                front,
                semantic,
                "--content-preservation-status",
                "PASS",
                "--journal-status",
                "NOT_APPLICABLE",
                "--render-status",
                "PASS",
                "--json",
            )
            self.assertEqual(passed.returncode, 0, passed.stdout + passed.stderr)
            self.assertEqual(json.loads(passed.stdout)["status"], "FORMAT_RELEASE_PASS")

            blocked = run_script(
                "validate_format_release.py",
                structural,
                front,
                semantic,
                "--content-preservation-status",
                "PASS",
                "--journal-status",
                "NOT_APPLICABLE",
                "--render-status",
                "NOT_ASSESSABLE",
                "--json",
            )
            self.assertEqual(blocked.returncode, 1)
            self.assertEqual(
                json.loads(blocked.stdout)["status"],
                "FORMAT_RELEASE_NOT_ASSESSABLE",
            )


if __name__ == "__main__":
    unittest.main()
