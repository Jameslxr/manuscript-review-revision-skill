#!/usr/bin/env python3
"""Regression tests for deterministic submission-package formatting."""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"


def run_script(name: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        text=True,
        capture_output=True,
        check=False,
    )


def build_mixed_cover_letter(path: Path) -> list[str]:
    document = Document()
    salutation = document.styles.add_style("Letter Salutation", WD_STYLE_TYPE.PARAGRAPH)
    salutation.font.name = "Arial"
    salutation.font.size = Pt(9)
    salutation.font.bold = True
    salutation.paragraph_format.space_after = Pt(8)
    salutation.paragraph_format.line_spacing = 1.0
    body = document.styles.add_style("Letter Body", WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after = Pt(10)
    body.paragraph_format.line_spacing = 1.15
    closing = document.styles.add_style("Letter Closing", WD_STYLE_TYPE.PARAGRAPH)
    closing.font.name = "Arial"
    closing.font.size = Pt(10)
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.line_spacing = 2.0
    document.add_paragraph("Dear Editors,", style=salutation)
    document.add_paragraph("We submit our review manuscript.", style=body)
    document.add_paragraph("It is suited to the journal readership.", style=body)
    document.add_paragraph("Sincerely,", style=closing)
    document.add_paragraph("James Li", style=body)
    document.add_paragraph("Department of Example Medicine", style=body)
    original = [p.text for p in document.paragraphs if p.text]
    document.save(path)
    return original


def apply_and_number(source: Path, normalized: Path, release: Path) -> None:
    applied = run_script(
        "apply_submission_package_profile.py",
        source,
        "--out",
        normalized,
        "--artifact-type",
        "cover-letter",
        "--line-spacing",
        "1.5",
    )
    if applied.returncode != 0:
        raise AssertionError(applied.stdout + applied.stderr)
    numbered = run_script(
        "enforce_docx_line_page_numbers.py", normalized, "--out", release
    )
    if numbered.returncode != 0:
        raise AssertionError(numbered.stdout + numbered.stderr)


class SubmissionPackageTests(unittest.TestCase):
    def test_explicit_cover_roles_survive_leading_blank_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "output.docx"
            document = Document()
            document.add_paragraph("")
            document.add_paragraph("Dear Editorial Office,")
            document.add_paragraph("Please consider this manuscript.")
            document.add_paragraph("Sincerely,")
            document.add_paragraph("James Li")
            document.save(source)
            result = run_script(
                "apply_submission_package_profile.py",
                source,
                "--out",
                output,
                "--artifact-type",
                "cover-letter",
                "--salutation-paragraph",
                "2",
                "--closing-paragraph",
                "4",
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertEqual(Document(output).paragraphs[0].text, "Dear Editorial Office,")

    def test_cover_profile_repairs_all_roles_and_signature_rhythm(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            report_path = root / "package.json"
            original = build_mixed_cover_letter(source)
            apply_and_number(source, normalized, release)

            audited = run_script(
                "audit_docx_submission_package.py",
                release,
                "--artifact-type",
                "cover-letter",
                "--expected-line-spacing",
                "1.5",
                "--output-json",
                report_path,
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "SUBMISSION_PACKAGE_PASS")
            self.assertEqual(report["roles"]["signature_paragraph_count"], 2)

            document = Document(normalized)
            self.assertEqual(original, [p.text for p in document.paragraphs if p.text])
            texts = [p.text for p in document.paragraphs]
            self.assertEqual(
                texts,
                [
                    "Dear Editors,",
                    "",
                    "We submit our review manuscript.",
                    "",
                    "It is suited to the journal readership.",
                    "",
                    "Sincerely,",
                    "",
                    "James Li",
                    "Department of Example Medicine",
                ],
            )
            for paragraph in document.paragraphs:
                self.assertAlmostEqual(float(paragraph.paragraph_format.line_spacing), 1.5)
                self.assertEqual(paragraph.paragraph_format.space_before.pt, 0.0)
                self.assertEqual(paragraph.paragraph_format.space_after.pt, 0.0)

    def test_cover_audit_rejects_mixed_line_spacing_and_missing_blank(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            broken = root / "broken.docx"
            build_mixed_cover_letter(source)
            apply_and_number(source, normalized, release)
            document = Document(release)
            body = next(p for p in document.paragraphs if p.text.startswith("We submit"))
            body.paragraph_format.line_spacing = 1.0
            blank = document.paragraphs[1]
            blank._element.getparent().remove(blank._element)
            document.save(broken)

            audited = run_script(
                "audit_docx_submission_package.py",
                broken,
                "--artifact-type",
                "cover-letter",
                "--expected-line-spacing",
                "1.5",
            )
            self.assertEqual(audited.returncode, 1)
            self.assertIn("PACKAGE_LINE_SPACING_MISMATCH", audited.stdout)
            self.assertIn("PACKAGE_BLANK_BOUNDARY_MISMATCH", audited.stdout)

    def test_generic_profile_keeps_list_items_compact(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            normalized = root / "normalized.docx"
            release = root / "release.docx"
            document = Document()
            document.add_paragraph("Highlights")
            first = document.add_paragraph("First highlight")
            second = document.add_paragraph("Second highlight")
            for paragraph in (first, second):
                num_pr = OxmlElement("w:numPr")
                paragraph._p.get_or_add_pPr().append(num_pr)
            document.add_paragraph("Declaration text")
            document.save(source)

            applied = run_script(
                "apply_submission_package_profile.py",
                source,
                "--out",
                normalized,
                "--artifact-type",
                "generic",
            )
            self.assertEqual(applied.returncode, 0, applied.stdout + applied.stderr)
            numbered = run_script(
                "enforce_docx_line_page_numbers.py", normalized, "--out", release
            )
            self.assertEqual(numbered.returncode, 0, numbered.stdout + numbered.stderr)
            audited = run_script(
                "audit_docx_submission_package.py",
                release,
                "--artifact-type",
                "generic",
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            texts = [p.text for p in Document(normalized).paragraphs]
            self.assertEqual(
                texts,
                ["Highlights", "", "First highlight", "Second highlight", "", "Declaration text"],
            )

    def test_package_release_is_fail_closed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            report = root / "package.json"
            report.write_text(
                json.dumps({"status": "SUBMISSION_PACKAGE_PASS"}), encoding="utf-8"
            )
            passed = run_script(
                "validate_submission_package_release.py",
                report,
                "--content-preservation-status",
                "PASS",
                "--journal-status",
                "NOT_APPLICABLE",
                "--render-status",
                "PASS",
            )
            self.assertEqual(passed.returncode, 0, passed.stdout + passed.stderr)
            self.assertIn("PACKAGE_FORMAT_RELEASE_PASS", passed.stdout)
            blocked = run_script(
                "validate_submission_package_release.py",
                report,
                "--content-preservation-status",
                "PASS",
                "--journal-status",
                "NOT_APPLICABLE",
                "--render-status",
                "NOT_ASSESSABLE",
            )
            self.assertEqual(blocked.returncode, 1)
            self.assertIn("PACKAGE_FORMAT_RELEASE_NOT_ASSESSABLE", blocked.stdout)


if __name__ == "__main__":
    unittest.main()
