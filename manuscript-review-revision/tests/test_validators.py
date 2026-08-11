#!/usr/bin/env python3
"""Representative tests for manuscript-review-revision validators."""

from __future__ import annotations

import csv
import hashlib
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SKILL_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = SKILL_ROOT / "scripts"
CORE_ROLES = [
    "journal-priority",
    "domain-science",
    "study-design",
    "statistics-reproducibility",
    "claim-evidence-reference",
]
CONCERN_COLUMNS = [
    "concern_id",
    "issue_key",
    "reviewer_id",
    "role_id",
    "axis",
    "role_scope",
    "severity",
    "claim_pointer",
    "evidence_pointer",
    "evidence_status",
    "concern",
    "finding_class",
    "defensibility_after_claim_narrowing",
    "resolution_mode",
    "resolution_test",
    "journal_gate",
    "confidence",
    "consensus_status",
    "disposition",
]
AXIS_OWNERS = {
    "journal-fit": "journal-priority",
    "novelty-significance": "journal-priority",
    "mechanism-evidence": "domain-science",
    "experimental-design": "study-design",
    "statistical-rigor": "statistics-reproducibility",
    "reproducibility": "statistics-reproducibility",
    "clinical-validity": "domain-science",
    "ethical-governance": "study-design",
    "data-resource-quality": "statistics-reproducibility",
    "figures-and-tables": "claim-evidence-reference",
    "writing-clarity": "claim-evidence-reference",
    "claim-moderation": "claim-evidence-reference",
    "causal-vs-correlative": "study-design",
    "reference-support": "claim-evidence-reference",
}
FORMAT_CHECK_CATEGORIES = [
    "article-type",
    "file-format",
    "title-page",
    "anonymization",
    "abstract",
    "main-text",
    "section-order",
    "references",
    "figures",
    "tables",
    "supplements",
    "line-numbering",
    "page-numbering",
    "statistics",
    "reporting-guidelines",
    "ethics-registration",
    "data-code",
    "declarations",
    "cover-letter",
    "submission-files",
]


def run_script(name: str, *args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPTS / name), *(str(arg) for arg in args)],
        text=True,
        capture_output=True,
        check=False,
    )


def file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def add_dynamic_page_field(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, display, end):
        run._r.append(node)


def apply_required_docx_numbering(
    document: Document,
    *,
    line_numbers: bool = True,
    page_numbers: bool = True,
) -> None:
    for section in document.sections:
        if line_numbers:
            line_numbers_node = OxmlElement("w:lnNumType")
            line_numbers_node.set(qn("w:countBy"), "1")
            line_numbers_node.set(qn("w:restart"), "continuous")
            section._sectPr.append(line_numbers_node)
        if page_numbers:
            add_dynamic_page_field(section.footer.paragraphs[0])


def write_panel(
    root: Path,
    *,
    duplicate_task: bool = False,
    corrupt_report_hash: bool = False,
    oversized_first_report: bool = False,
) -> Path:
    frozen_hashes = {
        "manuscript_sha256": "a" * 64,
        "journal_profile_sha256": "b" * 64,
        "shared_fact_base_sha256": "c" * 64,
    }
    reviewers = []
    for index, role in enumerate(CORE_ROLES, start=1):
        report = root / f"reviewer_{index:02d}.md"
        report_text = f"# Reviewer {index}\n\nEvidence-based report for {role}.\n"
        if oversized_first_report and index == 1:
            report_text += " ".join(["finding"] * 1801)
        report.write_text(report_text, encoding="utf-8")
        task_number = 1 if duplicate_task and index == 2 else index
        reviewers.append(
            {
                "agent_id": f"agent-{index}",
                "host_task_id": f"host-task-{task_number}",
                "receipt_source": "HOST_NATIVE",
                "context_mode": "FRESH_NON_FORK",
                "role_id": role,
                "role": role.replace("-", " "),
                "seat_type": "CORE",
                "primary_axes": [
                    axis for axis, owner in AXIS_OWNERS.items() if owner == role
                ],
                "independent": True,
                "saw_other_reviews": False,
                "status": "COMPLETED",
                "started_at": f"2026-07-19T10:0{index}:00-05:00",
                "completed_at": f"2026-07-19T10:1{index}:00-05:00",
                "input_hashes": frozen_hashes,
                "report_path": report.name,
                "report_sha256": (
                    "d" * 64
                    if corrupt_report_hash and index == 1
                    else file_sha256(report)
                ),
            }
        )
    panel = root / "panel.json"
    panel.write_text(
        json.dumps(
            {
                "panel_schema_version": "2.1",
                "skill_version": "1.5.0",
                "host": "Test Host",
                "host_version": "1.0",
                "target_journal": "Example Journal",
                "article_type": "Original Article",
                **frozen_hashes,
                "execution_mode": "independent_agents",
                "root_is_reviewer": False,
                "synthesis_started_before_reviews_completed": False,
                "review_policy": {
                    "core_reviewer_count": 5,
                    "maximum_panel_size": 6,
                    "max_concerns_per_reviewer": 8,
                    "max_blocking_major_per_reviewer": 6,
                    "max_minor_editorial_per_reviewer": 2,
                    "max_report_words": 1800,
                    "out_of_role_reporting": "BLOCKING_ONLY",
                    "overlap_target": 0.35,
                    "optional_seat_trigger": "NONE",
                    "axis_owners": AXIS_OWNERS,
                },
                "reviewers": reviewers,
            }
        ),
        encoding="utf-8",
    )
    return panel


def write_tolerance_card(
    path: Path,
    *,
    exact_matches: int = 5,
    substitution_reason: str = "",
) -> None:
    comparators = []
    for index in range(1, 6):
        comparators.append(
            {
                "id": f"P{index:02d}",
                "citation": f"Accepted comparator paper {index}",
                "url": f"https://journal.example.org/article-{index}",
                "publication_date": f"2026-0{index}-01",
                "journal": "Example Journal",
                "article_type": "Original Article",
                "match_level": (
                    "EXACT_JOURNAL_AND_TYPE"
                    if index <= exact_matches
                    else "EXACT_JOURNAL_ADJACENT_TYPE"
                ),
                "design": "Retrospective cohort",
                "unit_of_inference": "Patient",
                "scale": "100 patients",
                "validation": "Internal validation",
                "limitations": ["Single-center design"],
                "claim_ceiling": "Association only",
                "data_code_access": "Controlled access",
                "comparability_notes": "Same disease area and outcome family",
            }
        )
    path.write_text(
        json.dumps(
            {
                "target_journal": "Example Journal",
                "article_type": "Original Article",
                "accessed_at": "2026-07-29",
                "status": "PASS",
                "comparators": comparators,
                "substitution_reason": substitution_reason,
                "synthesis": {
                    "mandatory_official_requirements": ["Ethics statement"],
                    "validity_floor": ["Correct unit of analysis"],
                    "competitiveness_expectations": ["Field-level contribution"],
                    "accepted_with_limitations": ["Disclosed single-center design"],
                    "optional_strengthening": ["Additional external cohort"],
                    "calibration_boundary": (
                        "Published comparators calibrate scope, not scientific truth."
                    ),
                },
            }
        ),
        encoding="utf-8",
    )


def write_format_plan(
    path: Path,
    *,
    drop_category: str | None = None,
    paragraph_separation: str = "literal-blank",
    paragraph_separation_basis: str = "USER_GLOBAL_INVARIANT",
    space_after_pt: float = 0,
    line_numbering: str = "continuous",
    line_numbering_basis: str = "USER_GLOBAL_INVARIANT",
    page_numbering: str = "continuous",
    page_numbering_basis: str = "USER_GLOBAL_INVARIANT",
    unresolved_category: str | None = None,
    plan_status: str = "PASS",
) -> None:
    source_url = "https://journal.example.org/authors"
    checks = []
    for index, category in enumerate(FORMAT_CHECK_CATEGORIES, start=1):
        if category == drop_category:
            continue
        is_global_numbering = category in {"line-numbering", "page-numbering"}
        checks.append(
            {
                "id": f"F{index:03d}",
                "category": category,
                "requirement": f"Resolve the {category} rule for this submission.",
                "implementation": f"Apply the resolved {category} rule to the package.",
                "verification": f"Inspect the output for the {category} requirement.",
                "deliverable": "manuscript_clean.docx",
                "basis": (
                    "USER_GLOBAL_INVARIANT"
                    if is_global_numbering
                    else "OFFICIAL_GUIDE"
                ),
                "source_url": None if is_global_numbering else source_url,
                "mandatory": True,
                "status": (
                    "NOT_ASSESSABLE"
                    if category == unresolved_category
                    else "RESOLVED"
                ),
            }
        )

    path.write_text(
        json.dumps(
            {
                "schema_version": "1.1",
                "target_journal": "Example Journal",
                "article_type": "Original Article",
                "submission_stage": "initial",
                "accessed_at": "2026-08-10",
                "plan_status": plan_status,
                "journal_profile_sha256": "a" * 64,
                "official_sources": [
                    {
                        "title": "Author guide",
                        "url": source_url,
                        "accessed_at": "2026-08-10",
                        "official": True,
                    }
                ],
                "style_contract": {
                    "paragraph_separation": paragraph_separation,
                    "paragraph_separation_basis": paragraph_separation_basis,
                    "line_spacing": "double",
                    "line_spacing_basis": "CONSERVATIVE_FALLBACK",
                    "line_spacing_rule_strength": "UNSPECIFIED",
                    "line_spacing_source_excerpt": (
                        "No binding line-spacing value appears in the current guide."
                    ),
                    "line_numbering": line_numbering,
                    "line_numbering_basis": line_numbering_basis,
                    "page_numbering": page_numbering,
                    "page_numbering_basis": page_numbering_basis,
                    "page_number_position": "upper-right",
                    "page_number_position_basis": "CONSERVATIVE_FALLBACK",
                    "front_matter_alignment": "left",
                    "front_matter_alignment_basis": "CONSERVATIVE_FALLBACK",
                    "anonymization_mode": "unblinded",
                    "body_font_family": "Times New Roman",
                    "body_font_size_pt": 12,
                    "title_font_size_pt": 15,
                    "table_font_size_pt": 12,
                    "font_basis": "CONSERVATIVE_FALLBACK",
                    "font_rule_strength": "UNSPECIFIED",
                    "font_source_excerpt": (
                        "No binding font value appears in the current guide."
                    ),
                    "text_color_hex": "000000",
                    "space_before_pt": 0,
                    "space_after_pt": space_after_pt,
                    "body_styles": ["Normal", "Body Text"],
                    "page_size": "US Letter",
                    "margins": "1 inch on all sides",
                    "columns": 1,
                    "source_urls": [source_url],
                },
                "checks": checks,
            }
        ),
        encoding="utf-8",
    )


def write_format_audit(
    plan_path: Path,
    audit_path: Path,
    *,
    drop_check_id: str | None = None,
    mechanical_status: str = "PASS",
    visual_status: str = "PASS",
    front_matter_status: str = "PASS",
    content_preservation_status: str = "PASS",
    format_release_status: str = "FORMAT_RELEASE_PASS",
    page_number_position: str | None = None,
    overall_status: str = "PASS",
) -> None:
    plan = json.loads(plan_path.read_text(encoding="utf-8"))
    manuscript_path = audit_path.parent / "manuscript_clean.docx"
    manuscript = Document()
    manuscript.add_paragraph("Validated manuscript output.")
    manuscript.save(manuscript_path)

    checks = []
    for check in plan["checks"]:
        if check["id"] == drop_check_id:
            continue
        checks.append(
            {
                "check_id": check["id"],
                "status": "PASS",
                "evidence": f"Inspected {check['category']} in the output package.",
                "output": str(manuscript_path),
            }
        )

    audit_path.write_text(
        json.dumps(
            {
                "schema_version": "1.0",
                "target_journal": plan["target_journal"],
                "article_type": plan["article_type"],
                "submission_stage": plan["submission_stage"],
                "plan_sha256": file_sha256(plan_path),
                "overall_status": overall_status,
                "manuscripts": [
                    {
                        "role": "clean",
                        "path": str(manuscript_path),
                        "sha256": file_sha256(manuscript_path),
                        "mechanical_status": mechanical_status,
                        "paragraph_separation": plan["style_contract"][
                            "paragraph_separation"
                        ],
                        "line_spacing": plan["style_contract"]["line_spacing"],
                        "line_numbering": plan["style_contract"][
                            "line_numbering"
                        ],
                        "page_numbering": plan["style_contract"][
                            "page_numbering"
                        ],
                        "page_number_position": (
                            page_number_position
                            or plan["style_contract"]["page_number_position"]
                        ),
                        "front_matter_status": front_matter_status,
                        "front_matter_alignment": plan["style_contract"][
                            "front_matter_alignment"
                        ],
                        "anonymization_mode": plan["style_contract"][
                            "anonymization_mode"
                        ],
                        "content_preservation_status": content_preservation_status,
                        "format_release_status": format_release_status,
                        "issue_count": 0 if mechanical_status == "PASS" else 1,
                        "rendered_page_count": 2,
                        "inspected_pages": [1, 2],
                        "visual_status": visual_status,
                    }
                ],
                "checks": checks,
            }
        ),
        encoding="utf-8",
    )


def write_concern_ledger(path: Path, rows: list[dict[str, str]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=CONCERN_COLUMNS, delimiter="\t")
        writer.writeheader()
        writer.writerows(rows)


def concern_row(
    concern_id: str,
    issue_key: str,
    reviewer_number: int,
    role_id: str,
    *,
    consensus_status: str = "UNIQUE",
    evidence_status: str = "LOCATED",
    evidence_pointer: str = "Results > Validation, paragraph 2",
    axis: str = "experimental-design",
    role_scope: str = "PRIMARY",
    severity: str = "MAJOR",
    finding_class: str = "CORRECTABLE_BEFORE_SUBMISSION",
    defensibility: str | None = None,
    resolution_mode: str | None = None,
    disposition: str | None = None,
) -> dict[str, str]:
    if defensibility is None:
        defensibility = (
            "NOT_ASSESSABLE"
            if evidence_status == "NOT_ASSESSABLE"
            else "NOT_DEFENSIBLE"
            if severity == "BLOCKING"
            else "REMAINS_DEFENSIBLE"
        )
    if resolution_mode is None:
        resolution_mode = {
            "FATAL_VALIDITY_FLAW": "NO_DEFENSIBLE_REMEDY",
            "ACCEPTABLE_INHERENT_LIMITATION": "LIMITATION_DISCLOSURE",
            "OPTIONAL_STRENGTHENING": "NEW_ANALYSIS_OR_EXPERIMENT",
        }.get(finding_class, "CLAIM_NARROWING")
    if disposition is None:
        disposition = (
            "NOT_ASSESSABLE"
            if evidence_status == "NOT_ASSESSABLE"
            else "ACCEPTABLE_LIMITATION"
            if finding_class == "ACCEPTABLE_INHERENT_LIMITATION"
            else "OPEN"
        )
    return {
        "concern_id": concern_id,
        "issue_key": issue_key,
        "reviewer_id": f"agent-{reviewer_number}",
        "role_id": role_id,
        "axis": axis,
        "role_scope": role_scope,
        "severity": severity,
        "claim_pointer": "Discussion, paragraph 1: external validity",
        "evidence_pointer": evidence_pointer,
        "evidence_status": evidence_status,
        "concern": "The stated generalization exceeds the tested cohort.",
        "finding_class": finding_class,
        "defensibility_after_claim_narrowing": defensibility,
        "resolution_mode": resolution_mode,
        "resolution_test": "Narrow the claim or add an independent external cohort.",
        "journal_gate": "evidence threshold",
        "confidence": "0.9",
        "consensus_status": consensus_status,
        "disposition": disposition,
    }


class ValidatorTests(unittest.TestCase):
    def test_acceptance_tolerance_card_with_five_exact_matches_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            card = Path(temp) / "tolerance.json"
            write_tolerance_card(card)
            result = run_script("validate_acceptance_tolerance.py", card)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Acceptance-tolerance validation: PASS", result.stdout)
            self.assertIn("Exact journal/type matches: 5", result.stdout)

    def test_comparator_substitution_requires_disclosed_reason(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            card = Path(temp) / "tolerance.json"
            write_tolerance_card(card, exact_matches=4)
            result = run_script("validate_acceptance_tolerance.py", card)
            self.assertEqual(result.returncode, 1)
            self.assertIn("requires a non-empty substitution_reason", result.stdout)

    def test_valid_journal_profile_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "profile.json"
            source_url = "https://journal.example.org/authors"
            path.write_text(
                json.dumps(
                    {
                        "target_journal": "Example Journal",
                        "article_type": "Original Article",
                        "submission_stage": "initial",
                        "accessed_at": "2026-07-17",
                        "profile_status": "PASS",
                        "official_sources": [
                            {
                                "title": "Author guide",
                                "url": source_url,
                                "accessed_at": "2026-07-17",
                                "official": True,
                            }
                        ],
                        "requirements": [
                            {
                                "id": "J001",
                                "category": "format",
                                "text": "Use a standard manuscript file.",
                                "source_url": source_url,
                                "applies_to": "initial",
                                "mandatory": True,
                                "status": "MET",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            result = run_script("validate_journal_profile.py", path)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("validation: PASS", result.stdout)

    def test_pass_profile_with_pending_mandatory_rule_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "profile.json"
            source_url = "https://journal.example.org/authors"
            path.write_text(
                json.dumps(
                    {
                        "target_journal": "Example Journal",
                        "article_type": "Review",
                        "submission_stage": "initial",
                        "accessed_at": "2026-07-17",
                        "profile_status": "PASS",
                        "official_sources": [
                            {
                                "title": "Author guide",
                                "url": source_url,
                                "accessed_at": "2026-07-17",
                                "official": True,
                            }
                        ],
                        "requirements": [
                            {
                                "id": "J001",
                                "category": "references",
                                "text": "Reference limit must be confirmed.",
                                "source_url": source_url,
                                "applies_to": "initial",
                                "mandatory": True,
                                "status": "PENDING",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            result = run_script("validate_journal_profile.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("mandatory requirements are unresolved", result.stdout)

    def test_complete_journal_format_plan_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan.json"
            write_format_plan(path)
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Journal format plan validation: PASS", result.stdout)
            self.assertIn("Required categories covered: 20/20", result.stdout)

    def test_example_only_ten_point_font_cannot_override_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-example-font.json"
            write_format_plan(path)
            plan = json.loads(path.read_text(encoding="utf-8"))
            style = plan["style_contract"]
            style["body_font_size_pt"] = 10
            style["font_rule_strength"] = "EXAMPLE_ONLY"
            style["font_source_excerpt"] = (
                "Use a normal font, e.g., 10-point Times Roman."
            )
            path.write_text(json.dumps(plan), encoding="utf-8")
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("15/12/12 pt fallback", result.stdout)

    def test_explicit_twelve_point_and_one_point_five_rules_pass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-explicit-typography.json"
            write_format_plan(path)
            plan = json.loads(path.read_text(encoding="utf-8"))
            style = plan["style_contract"]
            style["font_basis"] = "OFFICIAL_GUIDE"
            style["font_rule_strength"] = "EXPLICIT_REQUIREMENT"
            style["font_source_excerpt"] = "Use 12-point Times New Roman."
            style["line_spacing"] = "1.5"
            style["line_spacing_basis"] = "OFFICIAL_GUIDE"
            style["line_spacing_rule_strength"] = "EXPLICIT_REQUIREMENT"
            style["line_spacing_source_excerpt"] = "Use 1.5 line spacing."
            path.write_text(json.dumps(plan), encoding="utf-8")
            result = run_script(
                "validate_journal_format_plan.py", path, "--require-pass"
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_example_only_typography_cannot_be_mandatory_profile_rule(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "profile-example-font.json"
            source_url = "https://journal.example.org/authors"
            path.write_text(
                json.dumps(
                    {
                        "target_journal": "Example Journal",
                        "article_type": "Original Article",
                        "submission_stage": "initial",
                        "accessed_at": "2026-08-11",
                        "profile_status": "PASS",
                        "official_sources": [
                            {
                                "title": "Author guide",
                                "url": source_url,
                                "accessed_at": "2026-08-11",
                                "official": True,
                            }
                        ],
                        "requirements": [
                            {
                                "id": "FONT-1",
                                "category": "font",
                                "text": "Example font wording inspected.",
                                "source_excerpt": (
                                    "Use a normal font, e.g., 10-point Times Roman."
                                ),
                                "rule_strength": "EXAMPLE_ONLY",
                                "source_url": source_url,
                                "applies_to": "initial manuscript",
                                "mandatory": True,
                                "status": "MET",
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            result = run_script("validate_journal_profile.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("cannot be mandatory", result.stdout)

    def test_journal_format_plan_requires_resolved_front_matter_contract(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-front-matter.json"
            write_format_plan(path)
            plan = json.loads(path.read_text(encoding="utf-8"))
            del plan["style_contract"]["front_matter_alignment"]
            path.write_text(json.dumps(plan), encoding="utf-8")
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("front_matter_alignment", result.stdout)

    def test_journal_format_plan_requires_all_categories(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-missing-cover-letter.json"
            write_format_plan(path, drop_category="cover-letter")
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "Missing required journal-format check categories: cover-letter",
                result.stdout,
            )

    def test_journal_template_paragraph_override_is_always_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-invalid-template-override.json"
            write_format_plan(
                path,
                paragraph_separation="journal-template",
                paragraph_separation_basis="OFFICIAL_TEMPLATE",
            )
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "paragraph_separation must be literal-blank",
                result.stdout,
            )

    def test_literal_blank_format_plan_requires_zero_paragraph_spacing(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-nonzero-spacing.json"
            write_format_plan(path, space_after_pt=6)
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "literal-blank mode requires style_contract.space_after_pt=0",
                result.stdout,
            )

    def test_format_plan_cannot_disable_continuous_line_or_page_numbering(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-numbering-disabled.json"
            write_format_plan(
                path,
                line_numbering="none",
                line_numbering_basis="OFFICIAL_TEMPLATE",
                page_numbering="per-section",
                page_numbering_basis="OFFICIAL_GUIDE",
            )
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "style_contract.line_numbering must be continuous",
                result.stdout,
            )
            self.assertIn(
                "style_contract.page_numbering must be continuous",
                result.stdout,
            )

    def test_pass_format_plan_cannot_hide_unresolved_mandatory_rule(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-unresolved.json"
            write_format_plan(path, unresolved_category="statistics")
            result = run_script("validate_journal_format_plan.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "plan_status is PASS but mandatory checks are unresolved",
                result.stdout,
            )

    def test_formatting_gate_rejects_draft_plan(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "format-plan-draft.json"
            write_format_plan(path, plan_status="DRAFT")
            result = run_script(
                "validate_journal_format_plan.py", path, "--require-pass"
            )
            self.assertEqual(result.returncode, 1)
            self.assertIn("--require-pass requires plan_status PASS", result.stdout)

    def test_complete_journal_format_audit_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            plan = root / "format-plan.json"
            audit = root / "format-audit.json"
            write_format_plan(plan)
            write_format_audit(plan, audit)
            result = run_script("validate_journal_format_audit.py", plan, audit)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Journal format audit validation: PASS", result.stdout)
            self.assertIn("Plan checks: 20", result.stdout)

    def test_journal_format_audit_requires_every_plan_check(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            plan = root / "format-plan.json"
            audit = root / "format-audit-missing-check.json"
            write_format_plan(plan)
            write_format_audit(plan, audit, drop_check_id="F019")
            result = run_script("validate_journal_format_audit.py", plan, audit)
            self.assertEqual(result.returncode, 1)
            self.assertIn("Audit is missing format-plan check IDs: F019", result.stdout)

    def test_journal_format_audit_pass_requires_mechanical_and_visual_pass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            plan = root / "format-plan.json"
            audit = root / "format-audit-failed-qa.json"
            write_format_plan(plan)
            write_format_audit(
                plan,
                audit,
                mechanical_status="FAIL",
                visual_status="FAIL",
                overall_status="PASS",
            )
            result = run_script("validate_journal_format_audit.py", plan, audit)
            self.assertEqual(result.returncode, 1)
            self.assertIn("mechanical PASS with zero issues", result.stdout)
            self.assertIn("visual_status PASS", result.stdout)

    def test_journal_format_audit_requires_front_matter_preservation_and_release(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            plan = root / "format-plan.json"
            audit = root / "format-audit-failed-release.json"
            write_format_plan(plan)
            write_format_audit(
                plan,
                audit,
                front_matter_status="FAIL",
                content_preservation_status="FAIL",
                format_release_status="FORMAT_RELEASE_FAIL",
                page_number_position="lower-center",
                overall_status="PASS",
            )
            result = run_script("validate_journal_format_audit.py", plan, audit)
            self.assertEqual(result.returncode, 1)
            self.assertIn("page_number_position does not match", result.stdout)
            self.assertIn("front_matter_status PASS", result.stdout)
            self.assertIn("content preservation PASS", result.stdout)
            self.assertIn("FORMAT_RELEASE_PASS", result.stdout)

    def test_journal_format_audit_rejects_nonpassing_plan(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            plan = root / "format-plan-draft.json"
            audit = root / "format-audit.json"
            write_format_plan(plan, plan_status="DRAFT")
            write_format_audit(plan, audit)
            result = run_script("validate_journal_format_audit.py", plan, audit)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "format plan must have plan_status PASS", result.stdout
            )

    def test_five_agent_panel_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Reviewer agents: 5", result.stdout)
            self.assertIn("Validated receipts: 5", result.stdout)

    def test_duplicate_task_id_and_inconsistent_input_fail(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root, duplicate_task=True)
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("Duplicate host_task_id", result.stdout)

            panel = write_panel(root)
            content = json.loads(panel.read_text(encoding="utf-8"))
            content["reviewers"][0]["input_hashes"]["manuscript_sha256"] = "f" * 64
            panel.write_text(json.dumps(content), encoding="utf-8")
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("does not match the frozen panel input", result.stdout)

    def test_report_hash_mismatch_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            panel = write_panel(Path(temp), corrupt_report_hash=True)
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("report_sha256 does not match", result.stdout)

    def test_not_assessable_reviewer_does_not_count_as_completed(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            panel = write_panel(Path(temp))
            content = json.loads(panel.read_text(encoding="utf-8"))
            content["reviewers"][0]["status"] = "NOT_ASSESSABLE"
            panel.write_text(json.dumps(content), encoding="utf-8")
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "Only 4 completed independent reviewer agents",
                result.stdout,
            )

    def test_concern_ledger_passes_with_valid_consensus(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            rows = [
                concern_row(
                    "C001",
                    "external-validation",
                    2,
                    "domain-science",
                    consensus_status="CONSENSUS",
                    role_scope="BLOCKING_CROSSOVER",
                    severity="BLOCKING",
                ),
                concern_row(
                    "C002",
                    "external-validation",
                    3,
                    "study-design",
                    consensus_status="CONSENSUS",
                    severity="BLOCKING",
                ),
                concern_row(
                    "C003",
                    "multiplicity",
                    4,
                    "statistics-reproducibility",
                    axis="statistical-rigor",
                ),
            ]
            write_concern_ledger(ledger, rows)
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Concern ledger validation: PASS", result.stdout)
            self.assertIn("Consensus / disagreement / unique: 1 / 0 / 1", result.stdout)

    def test_single_reviewer_cannot_claim_consensus(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "external-validation",
                        3,
                        "study-design",
                        consensus_status="CONSENSUS",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("fewer than two reviewers", result.stdout)

    def test_located_evidence_requires_pointer(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "external-validation",
                        3,
                        "study-design",
                        evidence_pointer="",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("requires a specific evidence_pointer", result.stdout)

    def test_acceptable_inherent_limitation_is_not_blocking(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "single-center-generalizability",
                        3,
                        "study-design",
                        finding_class="ACCEPTABLE_INHERENT_LIMITATION",
                        severity="MAJOR",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Concern ledger validation: PASS", result.stdout)

    def test_blocking_requires_failure_after_claim_narrowing(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "overstated-claim",
                        3,
                        "study-design",
                        severity="BLOCKING",
                        defensibility="REMAINS_DEFENSIBLE",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "BLOCKING requires NOT_DEFENSIBLE after claim narrowing",
                result.stdout,
            )

    def test_optional_strengthening_cannot_be_blocking(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "extra-validation-cohort",
                        3,
                        "study-design",
                        severity="BLOCKING",
                        finding_class="OPTIONAL_STRENGTHENING",
                        defensibility="NOT_DEFENSIBLE",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("OPTIONAL_STRENGTHENING", result.stdout)

    def test_blocking_requires_located_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "suspected-invalidity",
                        3,
                        "study-design",
                        severity="BLOCKING",
                        evidence_status="LOCATION_NOT_PROVIDED",
                        evidence_pointer="",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BLOCKING requires located manuscript evidence", result.stdout)

    def test_high_overlap_warns_without_inventing_diversity(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            rows = []
            for issue_number in range(1, 4):
                issue_key = f"shared-{issue_number}"
                rows.extend(
                    [
                        concern_row(
                            f"C{issue_number}A",
                            issue_key,
                            2,
                            "domain-science",
                            consensus_status="CONSENSUS",
                            role_scope="BLOCKING_CROSSOVER",
                            severity="BLOCKING",
                        ),
                        concern_row(
                            f"C{issue_number}B",
                            issue_key,
                            3,
                            "study-design",
                            consensus_status="CONSENSUS",
                            severity="BLOCKING",
                        ),
                    ]
                )
            write_concern_ledger(ledger, rows)
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("WARNING: Reviewer pair agent-2/agent-3", result.stdout)

    def test_seventh_reviewer_exceeds_panel_maximum(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            content = json.loads(panel.read_text(encoding="utf-8"))
            for index, role in ((6, "figure-narrative-reporting"), (7, "adversarial-review")):
                report = root / f"reviewer_{index:02d}.md"
                report.write_text(f"# Reviewer {index}\n", encoding="utf-8")
                receipt = dict(content["reviewers"][0])
                receipt.update(
                    {
                        "agent_id": f"agent-{index}",
                        "host_task_id": f"host-task-{index}",
                        "role_id": role,
                        "role": role.replace("-", " "),
                        "seat_type": "OPTIONAL",
                        "primary_axes": [],
                        "report_path": report.name,
                        "report_sha256": file_sha256(report),
                    }
                )
                content["reviewers"].append(receipt)
            content["review_policy"]["optional_seat_trigger"] = "multiple risks"
            panel.write_text(json.dumps(content), encoding="utf-8")
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("maximum is 6", result.stdout)
            self.assertIn("at most one OPTIONAL seat", result.stdout)

    def test_required_core_role_cannot_be_optional(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            panel = write_panel(Path(temp))
            content = json.loads(panel.read_text(encoding="utf-8"))
            content["reviewers"][0]["seat_type"] = "OPTIONAL"
            content["review_policy"]["optional_seat_trigger"] = "editorial risk"
            panel.write_text(json.dumps(content), encoding="utf-8")
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "Required core role 'journal-priority' must use seat_type CORE",
                result.stdout,
            )

    def test_reviewer_report_budget_is_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            panel = write_panel(Path(temp), oversized_first_report=True)
            result = run_script("validate_review_panel.py", panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("word-equivalent units; limit is 1800", result.stdout)

    def test_concern_budget_is_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            rows = [
                concern_row(
                    f"C{index:03d}",
                    f"design-issue-{index}",
                    3,
                    "study-design",
                )
                for index in range(1, 10)
            ]
            write_concern_ledger(ledger, rows)
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("has 9 concerns; limit is 8", result.stdout)

    def test_blocking_major_budget_is_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            rows = [
                concern_row(
                    f"C{index:03d}",
                    f"design-major-{index}",
                    3,
                    "study-design",
                    severity="MAJOR",
                )
                for index in range(1, 8)
            ]
            write_concern_ledger(ledger, rows)
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("has 7 BLOCKING/MAJOR concerns; limit is 6", result.stdout)

    def test_minor_editorial_budget_is_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            rows = [
                concern_row(
                    f"C{index:03d}",
                    f"design-minor-{index}",
                    3,
                    "study-design",
                    severity="MINOR",
                )
                for index in range(1, 4)
            ]
            write_concern_ledger(ledger, rows)
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn("has 3 MINOR/EDITORIAL concerns; limit is 2", result.stdout)

    def test_non_blocking_out_of_role_concern_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            panel = write_panel(root)
            ledger = root / "concerns.tsv"
            write_concern_ledger(
                ledger,
                [
                    concern_row(
                        "C001",
                        "design-issue",
                        2,
                        "domain-science",
                        role_scope="BLOCKING_CROSSOVER",
                        severity="MAJOR",
                    )
                ],
            )
            result = run_script("validate_concern_ledger.py", ledger, panel)
            self.assertEqual(result.returncode, 1)
            self.assertIn(
                "BLOCKING_CROSSOVER concerns must have BLOCKING severity",
                result.stdout,
            )

    def test_bounded_single_posture_verdict_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            verdict = Path(temp) / "verdict.md"
            verdict.write_text(
                "# Verdict\n\n`MAJOR_SCIENTIFIC_REWORK_REQUIRED`\n\n"
                "The central claim requires a leakage-free validation cohort.\n",
                encoding="utf-8",
            )
            result = run_script("validate_review_verdict.py", verdict)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Review verdict validation: PASS", result.stdout)

    def test_overlong_or_multiple_posture_verdict_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            verdict = Path(temp) / "verdict.md"
            verdict.write_text(
                "MAJOR_SCIENTIFIC_REWORK_REQUIRED RETARGET_RECOMMENDED\n"
                + " ".join(["finding"] * 901),
                encoding="utf-8",
            )
            result = run_script("validate_review_verdict.py", verdict)
            self.assertEqual(result.returncode, 1)
            self.assertIn("maximum is 900", result.stdout)
            self.assertIn("exactly one allowed review posture", result.stdout)

    def test_reference_direct_support_passes_only_with_full_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            ledger = Path(temp) / "references.tsv"
            columns = [
                "sentence_id",
                "atomic_claim",
                "claim_tier",
                "citation_key",
                "identifier",
                "metadata_status",
                "integrity_status",
                "evidence_basis",
                "support_grade",
                "placement_status",
                "format_status",
                "action",
            ]
            with ledger.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=columns, delimiter="\t")
                writer.writeheader()
                writer.writerow(
                    {
                        "sentence_id": "S001",
                        "atomic_claim": "The tested relationship was observed in HCC.",
                        "claim_tier": "A_MATERIAL",
                        "citation_key": "REF001",
                        "identifier": "doi:10.1000/example",
                        "metadata_status": "VERIFIED",
                        "integrity_status": "CLEAR",
                        "evidence_basis": "RESULTS_SECTION",
                        "support_grade": "DIRECT_SUPPORT",
                        "placement_status": "PRECISE",
                        "format_status": "PASS",
                        "action": "retain",
                    }
                )
            result = run_script("validate_reference_audit.py", ledger)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("validation: PASS", result.stdout)

    def test_metadata_only_direct_support_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            ledger = Path(temp) / "references.tsv"
            columns = [
                "sentence_id",
                "atomic_claim",
                "claim_tier",
                "citation_key",
                "identifier",
                "metadata_status",
                "integrity_status",
                "evidence_basis",
                "support_grade",
                "placement_status",
                "format_status",
                "action",
            ]
            with ledger.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=columns, delimiter="\t")
                writer.writeheader()
                writer.writerow(
                    {
                        "sentence_id": "S001",
                        "atomic_claim": "A causal mechanism is established.",
                        "claim_tier": "A_MATERIAL",
                        "citation_key": "REF001",
                        "identifier": "doi:10.1000/example",
                        "metadata_status": "VERIFIED",
                        "integrity_status": "CLEAR",
                        "evidence_basis": "METADATA_ONLY",
                        "support_grade": "DIRECT_SUPPORT",
                        "placement_status": "PRECISE",
                        "format_status": "PASS",
                        "action": "inspect full text",
                    }
                )
            result = run_script("validate_reference_audit.py", ledger)
            self.assertEqual(result.returncode, 1)
            self.assertIn("requires FULL_TEXT or RESULTS_SECTION", result.stdout)

    def test_unassessed_context_citation_is_advisory(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            ledger = Path(temp) / "references.tsv"
            columns = [
                "sentence_id",
                "atomic_claim",
                "claim_tier",
                "citation_key",
                "identifier",
                "metadata_status",
                "integrity_status",
                "evidence_basis",
                "support_grade",
                "placement_status",
                "format_status",
                "action",
            ]
            with ledger.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=columns, delimiter="\t")
                writer.writeheader()
                writer.writerow(
                    {
                        "sentence_id": "S001",
                        "atomic_claim": "HCC is a major form of primary liver cancer.",
                        "claim_tier": "C_CONTEXT",
                        "citation_key": "REF001",
                        "identifier": "",
                        "metadata_status": "UNVERIFIED",
                        "integrity_status": "NOT_CHECKED",
                        "evidence_basis": "UNAVAILABLE",
                        "support_grade": "NOT_ASSESSABLE",
                        "placement_status": "AMBIGUOUS",
                        "format_status": "NOT_ASSESSABLE",
                        "action": "sample for follow-up",
                    }
                )
            result = run_script("validate_reference_audit.py", ledger)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Reference audit validation: PASS", result.stdout)
            self.assertIn("Advisory rows: [2]", result.stdout)

    def test_unassessed_material_citation_blocks_pass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            ledger = Path(temp) / "references.tsv"
            columns = [
                "sentence_id",
                "atomic_claim",
                "claim_tier",
                "citation_key",
                "identifier",
                "metadata_status",
                "integrity_status",
                "evidence_basis",
                "support_grade",
                "placement_status",
                "format_status",
                "action",
            ]
            with ledger.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=columns, delimiter="\t")
                writer.writeheader()
                writer.writerow(
                    {
                        "sentence_id": "S001",
                        "atomic_claim": "The intervention improves overall survival.",
                        "claim_tier": "A_MATERIAL",
                        "citation_key": "REF001",
                        "identifier": "",
                        "metadata_status": "UNVERIFIED",
                        "integrity_status": "NOT_CHECKED",
                        "evidence_basis": "UNAVAILABLE",
                        "support_grade": "NOT_ASSESSABLE",
                        "placement_status": "AMBIGUOUS",
                        "format_status": "NOT_ASSESSABLE",
                        "action": "inspect the primary report",
                    }
                )
            result = run_script("validate_reference_audit.py", ledger)
            self.assertEqual(result.returncode, 1)
            self.assertIn("Reference audit validation: NOT_ASSESSABLE", result.stdout)
            self.assertIn("Blocked rows: [2]", result.stdout)

    def test_blue_heading_docx_fails_and_black_heading_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            blue_path = root / "blue.docx"
            blue = Document()
            title = blue.add_paragraph(style="Title")
            title_run = title.add_run("Blue manuscript title")
            title_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            heading = blue.add_heading("Introduction", level=1)
            heading.runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            blue.add_paragraph("Body text.")
            apply_required_docx_numbering(blue)
            blue.save(blue_path)

            blue_result = run_script("audit_docx_manuscript_style.py", blue_path)
            self.assertEqual(blue_result.returncode, 1)
            self.assertIn("NON_BLACK", blue_result.stdout)

            black_path = root / "black.docx"
            black = Document()
            black.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
            black.styles["Normal"].paragraph_format.line_spacing = 2.0
            black_title = black.add_paragraph()
            black_title_run = black_title.add_run("Black manuscript title")
            black_title_run.bold = True
            black_title_run.font.color.rgb = RGBColor(0, 0, 0)
            black.add_heading("Introduction", level=1)
            black.add_paragraph("Body text.")
            apply_required_docx_numbering(black)
            black.save(black_path)

            black_result = run_script("audit_docx_manuscript_style.py", black_path)
            self.assertEqual(
                black_result.returncode, 0, black_result.stdout + black_result.stderr
            )
            self.assertIn("MECHANICAL_PASS", black_result.stdout)

    def test_body_spacing_and_missing_literal_blank_fail_docx_audit(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "spaced.docx"
            document = Document()
            document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
            document.styles["Normal"].paragraph_format.space_before = Pt(0)
            document.styles["Normal"].paragraph_format.space_after = Pt(8)
            document.styles["Normal"].paragraph_format.line_spacing = 2.0
            document.add_heading("Introduction", level=1)
            document.add_paragraph("First body paragraph.")
            document.add_paragraph("Second body paragraph.")
            apply_required_docx_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BODY_PARAGRAPH_SPACE_AFTER", result.stdout)
            self.assertIn("MISSING_LITERAL_BLANK_PARAGRAPH", result.stdout)

    def test_real_empty_paragraph_with_zero_spacing_passes_docx_audit(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "literal-blank.docx"
            document = Document()
            document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
            document.styles["Normal"].paragraph_format.space_before = Pt(0)
            document.styles["Normal"].paragraph_format.space_after = Pt(0)
            document.styles["Normal"].paragraph_format.line_spacing = 2.0
            document.add_heading("Introduction", level=1)
            document.add_paragraph("First body paragraph.")
            separator = document.add_paragraph("")
            separator.paragraph_format.space_before = Pt(0)
            separator.paragraph_format.space_after = Pt(0)
            document.add_paragraph("Second body paragraph.")
            apply_required_docx_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Paragraph separation: literal-blank", result.stdout)
            self.assertIn("MECHANICAL_PASS", result.stdout)

    def test_docx_audit_rejects_journal_template_spacing_bypass(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "journal-template.docx"
            document = Document()
            document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
            document.styles["Normal"].paragraph_format.space_after = Pt(8)
            document.styles["Normal"].paragraph_format.line_spacing = 1.5
            document.add_heading("Introduction", level=1)
            document.add_paragraph("First body paragraph.")
            document.add_paragraph("Second body paragraph.")
            apply_required_docx_numbering(document)
            document.save(path)

            result = run_script(
                "audit_docx_manuscript_style.py",
                path,
                "--paragraph-separation",
                "journal-template",
                "--expected-line-spacing",
                "1.5",
            )
            self.assertEqual(result.returncode, 2)
            self.assertIn("invalid choice", result.stderr)

    def test_unclassified_custom_style_blocks_docx_delivery(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "unclassified-style.docx"
            document = Document()
            custom = document.styles.add_style(
                "Custom Front Matter", WD_STYLE_TYPE.PARAGRAPH
            )
            custom.base_style = None
            document.styles["Normal"].paragraph_format.space_before = Pt(0)
            document.styles["Normal"].paragraph_format.space_after = Pt(0)
            document.styles["Normal"].paragraph_format.line_spacing = 2.0
            document.add_paragraph("Custom title block", style=custom)
            document.add_paragraph("Body text.")
            apply_required_docx_numbering(document)
            document.save(path)

            blocked = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(blocked.returncode, 1)
            self.assertIn(
                "UNCLASSIFIED_NONEMPTY_PARAGRAPH_STYLE", blocked.stdout
            )

            classified = run_script(
                "audit_docx_manuscript_style.py",
                path,
                "--exclude-style",
                "Custom Front Matter",
            )
            self.assertEqual(
                classified.returncode,
                0,
                classified.stdout + classified.stderr,
            )

    def test_custom_body_style_must_be_explicitly_audited(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "custom-body-style.docx"
            document = Document()
            custom = document.styles.add_style(
                "Custom Manuscript Prose", WD_STYLE_TYPE.PARAGRAPH
            )
            custom.base_style = None
            custom.paragraph_format.space_before = Pt(0)
            custom.paragraph_format.space_after = Pt(0)
            custom.paragraph_format.line_spacing = 2.0
            document.add_paragraph("First body paragraph.", style=custom)
            document.add_paragraph("", style=custom)
            document.add_paragraph("Second body paragraph.", style=custom)
            apply_required_docx_numbering(document)
            document.save(path)

            blocked = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(blocked.returncode, 1)
            self.assertIn(
                "UNCLASSIFIED_NONEMPTY_PARAGRAPH_STYLE", blocked.stdout
            )

            audited = run_script(
                "audit_docx_manuscript_style.py",
                path,
                "--body-style",
                "Custom Manuscript Prose",
            )
            self.assertEqual(audited.returncode, 0, audited.stdout + audited.stderr)
            self.assertIn("MECHANICAL_PASS", audited.stdout)

    def test_docx_audit_cannot_disable_line_spacing_gate(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "line-spacing-off.docx"
            document = Document()
            document.add_paragraph("Body text.")
            apply_required_docx_numbering(document)
            document.save(path)

            result = run_script(
                "audit_docx_manuscript_style.py",
                path,
                "--expected-line-spacing",
                "off",
            )
            self.assertEqual(result.returncode, 2)
            self.assertIn("Unable to audit DOCX", result.stderr)

    def test_missing_line_and_page_numbers_block_docx_delivery(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "missing-numbering.docx"
            document = Document()
            normal = document.styles["Normal"].paragraph_format
            normal.space_before = Pt(0)
            normal.space_after = Pt(0)
            normal.line_spacing = 2.0
            document.add_paragraph("Body text.")
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("SECTION_LINE_NUMBERING_MISSING", result.stdout)
            self.assertIn("PAGE_NUMBER_FIELD_MISSING", result.stdout)

    def test_number_restarts_and_line_suppression_block_docx_delivery(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "numbering-restarts.docx"
            document = Document()
            normal = document.styles["Normal"].paragraph_format
            normal.space_before = Pt(0)
            normal.space_after = Pt(0)
            normal.line_spacing = 2.0
            first = document.add_paragraph("First section body.")
            first._p.get_or_add_pPr().append(OxmlElement("w:suppressLineNumbers"))
            document.add_section(WD_SECTION.NEW_PAGE)
            document.add_paragraph("Second section body.")
            apply_required_docx_numbering(document)

            second_line_numbers = document.sections[1]._sectPr.xpath(
                "./w:lnNumType"
            )[0]
            second_line_numbers.set(qn("w:restart"), "newPage")
            page_numbers = OxmlElement("w:pgNumType")
            page_numbers.set(qn("w:start"), "1")
            document.sections[1]._sectPr.append(page_numbers)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("SECTION_LINE_NUMBERING_NOT_CONTINUOUS", result.stdout)
            self.assertIn("LINE_NUMBER_SUPPRESSION_PRESENT", result.stdout)
            self.assertIn("PAGE_NUMBERING_NOT_CONTINUOUS", result.stdout)

    def test_numbering_enforcer_covers_multisection_active_page_stories(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = root / "source.docx"
            output = root / "numbered.docx"
            document = Document()
            normal = document.styles["Normal"].paragraph_format
            normal.space_before = Pt(0)
            normal.space_after = Pt(0)
            normal.line_spacing = 2.0
            document.sections[0].different_first_page_header_footer = True
            document.add_paragraph("First section body.")
            document.add_section(WD_SECTION.NEW_PAGE)
            document.sections[1].different_first_page_header_footer = True
            document.add_paragraph("Second section body.")
            document.settings.element.append(OxmlElement("w:evenAndOddHeaders"))
            document.save(source)

            enforced = run_script(
                "enforce_docx_line_page_numbers.py",
                source,
                "--out",
                output,
            )
            self.assertEqual(
                enforced.returncode,
                0,
                enforced.stdout + enforced.stderr,
            )

            result = run_script("audit_docx_manuscript_style.py", output)
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("Continuous line-numbered sections: 2/2", result.stdout)
            self.assertIn("Dynamic PAGE fields in active page stories: 6/6", result.stdout)
            self.assertIn("MECHANICAL_PASS", result.stdout)

    def test_hidden_word_autospacing_blocks_docx_delivery(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "hidden-autospacing.docx"
            document = Document()
            normal = document.styles["Normal"]
            normal.paragraph_format.space_before = Pt(0)
            normal.paragraph_format.space_after = Pt(0)
            normal.paragraph_format.line_spacing = 2.0
            style_properties = normal.element.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:afterAutospacing"), "1")
            style_properties.append(spacing)
            document.add_paragraph("First body paragraph.")
            document.add_paragraph("")
            document.add_paragraph("Second body paragraph.")
            apply_required_docx_numbering(document)
            document.save(path)

            blocked = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(blocked.returncode, 1)
            self.assertIn("BODY_PARAGRAPH_AUTOSPACING_AFTER", blocked.stdout)

            spacing.set(qn("w:afterAutospacing"), "0")
            document.save(path)
            cleared = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(cleared.returncode, 0, cleared.stdout + cleared.stderr)

    def test_wrong_body_and_separator_line_spacing_fail_docx_audit(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "wrong-line-spacing.docx"
            document = Document()
            document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
            document.styles["Normal"].paragraph_format.space_before = Pt(0)
            document.styles["Normal"].paragraph_format.space_after = Pt(0)
            document.styles["Normal"].paragraph_format.line_spacing = 1.15
            document.add_heading("Introduction", level=1)
            document.add_paragraph("First body paragraph.")
            document.add_paragraph("")
            document.add_paragraph("Second body paragraph.")
            apply_required_docx_numbering(document)
            document.save(path)

            result = run_script("audit_docx_manuscript_style.py", path)
            self.assertEqual(result.returncode, 1)
            self.assertIn("BODY_LINE_SPACING_MISMATCH", result.stdout)
            self.assertIn("BLANK_PARAGRAPH_LINE_SPACING_MISMATCH", result.stdout)

    def test_exact_and_at_least_line_spacing_tokens_pass_docx_audit(self) -> None:
        cases = (
            ("exact", 24, WD_LINE_SPACING.EXACTLY, "exact:24pt"),
            ("at-least", 14, WD_LINE_SPACING.AT_LEAST, "at-least:14pt"),
        )
        with tempfile.TemporaryDirectory() as temp:
            for label, points, rule, expected in cases:
                with self.subTest(label=label):
                    path = Path(temp) / f"{label}.docx"
                    document = Document()
                    document.styles["Heading 1"].font.color.rgb = RGBColor(0, 0, 0)
                    normal = document.styles["Normal"].paragraph_format
                    normal.space_before = Pt(0)
                    normal.space_after = Pt(0)
                    normal.line_spacing = Pt(points)
                    normal.line_spacing_rule = rule
                    document.add_heading("Introduction", level=1)
                    document.add_paragraph("First body paragraph.")
                    document.add_paragraph("")
                    document.add_paragraph("Second body paragraph.")
                    apply_required_docx_numbering(document)
                    document.save(path)

                    result = run_script(
                        "audit_docx_manuscript_style.py",
                        path,
                        "--expected-line-spacing",
                        expected,
                    )
                    self.assertEqual(
                        result.returncode, 0, result.stdout + result.stderr
                    )
                    self.assertIn("MECHANICAL_PASS", result.stdout)


if __name__ == "__main__":
    unittest.main()
