#!/usr/bin/env python3
"""Audit manuscript-wide typography and semantic vertical rhythm."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - environment-dependent error path
    print(
        "python-docx is required. Use the bundled workspace Python runtime.",
        file=sys.stderr,
    )
    raise SystemExit(2) from exc

from docx_semantic_rules import (  # noqa: E402
    CREDIT_HEADING_RE,
    CREDIT_INLINE_RE,
    DECLARATION_HEADING_RE,
    DECLARATION_INLINE_RE,
    HEADING_STYLE_RE,
    KEYWORD_LABEL_RE,
    credit_role_labels,
)
from audit_docx_front_matter import paragraph_font_sizes_pt  # noqa: E402
from audit_docx_manuscript_style import (  # noqa: E402
    DEFAULT_BODY_STYLES,
    automatic_spacing_sources,
    effective_line_spacing,
    effective_spacing,
    iter_style_chain,
    line_spacing_matches,
    normalize_style_token,
    paragraph_is_list,
    paragraph_is_structurally_empty,
    parse_line_spacing_spec,
)


ROLE_STYLE_RE = {
    "title": re.compile(r"^(?:manuscript\s+)?title$", re.I),
    "author_note": re.compile(r"^(?:manuscript\s+)?author\s+note$", re.I),
    "orcid": re.compile(r"^(?:manuscript\s+)?(?:orcid|identifiers?)$", re.I),
    "authors": re.compile(r"^(?:manuscript\s+)?authors?$", re.I),
    "affiliation": re.compile(r"^(?:manuscript\s+)?affiliations?$", re.I),
    "correspondence": re.compile(
        r"^(?:manuscript\s+)?(?:correspondence|corresponding\s+author)$", re.I
    ),
    "reference": re.compile(
        r"^(?:manuscript\s+)?(?:reference|bibliograph(?:y|ic))$", re.I
    ),
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Audit global manuscript line spacing, body-size typography, Keywords "
            "label emphasis, and semantic blank-line boundaries."
        )
    )
    parser.add_argument("document", type=Path)
    parser.add_argument("--expected-line-spacing", default="double")
    parser.add_argument("--expected-font-name", default="Times New Roman")
    parser.add_argument("--expected-body-font-size", type=float, default=12.0)
    parser.add_argument("--expected-title-font-size", type=float, default=15.0)
    parser.add_argument(
        "--expected-table-font-size",
        type=float,
        default=10.0,
        help="Expected table-cell size; defaults to 10 pt.",
    )
    parser.add_argument(
        "--expected-table-line-spacing",
        default="single",
        help="Expected table-cell line spacing; defaults to single.",
    )
    parser.add_argument(
        "--expected-table-rule-scheme",
        choices=("three-line", "full-grid", "preserve-official"),
        default="three-line",
        help=(
            "Expected table presentation. The journal-neutral fallback is a "
            "three-line table with no vertical rules or shading."
        ),
    )
    parser.add_argument("--body-style", action="append", default=[])
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--output-json", type=Path)
    return parser.parse_args()


def style_tokens(paragraph: Any) -> set[str]:
    if paragraph.style is None:
        return set()
    return {
        normalize_style_token(paragraph.style.name),
        normalize_style_token(paragraph.style.style_id),
    }


def classify_role(paragraph: Any, body_styles: set[str]) -> str | None:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if not text:
        return "blank"
    for role, pattern in ROLE_STYLE_RE.items():
        if any(pattern.search(token) for token in style_tokens(paragraph)):
            return role
    if KEYWORD_LABEL_RE.match(paragraph.text) or re.search(
        r"keyword", style_name, re.I
    ):
        return "keywords"
    if CREDIT_HEADING_RE.fullmatch(paragraph.text):
        return "credit-heading"
    if re.search(r"credit.*entry", style_name, re.I):
        return "credit-entry"
    if (
        HEADING_STYLE_RE.match(style_name)
        or style_name.lower().startswith("heading")
        or DECLARATION_HEADING_RE.fullmatch(paragraph.text)
    ):
        return "heading"
    if paragraph_is_list(paragraph):
        return "list-item"
    if CREDIT_INLINE_RE.match(paragraph.text):
        return "credit-inline"
    if DECLARATION_INLINE_RE.match(paragraph.text):
        return "declaration-inline"
    if style_tokens(paragraph) & body_styles:
        return "body"
    return None


def effective_bold(run: Any, paragraph: Any) -> bool:
    if run.bold is not None:
        return bool(run.bold)
    if paragraph.style is not None:
        for style in iter_style_chain(paragraph.style):
            if style.font.bold is not None:
                return bool(style.font.bold)
    return False


def label_weight_issues(
    paragraph: Any,
    pattern: re.Pattern[str],
    *,
    prefix: str,
) -> list[dict[str, object]]:
    match = pattern.match(paragraph.text)
    if match is None:
        return [{"code": f"{prefix}_LABEL_MISSING"}]
    label_end = match.end(1)
    consumed = 0
    label_not_bold = False
    content_bold = False
    for run in paragraph.runs:
        start = consumed
        end = start + len(run.text)
        consumed = end
        bold = effective_bold(run, paragraph)
        if start < label_end and run.text[: max(0, label_end - start)].strip() and not bold:
            label_not_bold = True
        if end > label_end:
            content_start = max(0, label_end - start)
            if run.text[content_start:].strip() and bold:
                content_bold = True
    issues: list[dict[str, object]] = []
    if label_not_bold:
        issues.append({"code": f"{prefix}_LABEL_NOT_BOLD"})
    if content_bold:
        issues.append({"code": f"{prefix}_CONTENT_BOLD"})
    return issues


def adjacent_blank_count(document: Any, node: Any, direction: int) -> int:
    children = list(document.element.body)
    index = children.index(node) + direction
    count = 0
    while 0 <= index < len(children):
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            break
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            break
        count += 1
        index += direction
    return count


def previous_nonblank_role(
    document: Any,
    node: Any,
    roles: dict[Any, str | None],
) -> str | None:
    children = list(document.element.body)
    index = children.index(node) - 1
    while index >= 0:
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            return None
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            return roles.get(candidate)
        index -= 1
    return None


def next_nonblank_role(
    document: Any,
    node: Any,
    roles: dict[Any, str | None],
) -> str | None:
    children = list(document.element.body)
    index = children.index(node) + 1
    while index < len(children):
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            return None
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            return roles.get(candidate)
        index += 1
    return None


def visible_run_text(run: Any) -> str:
    return "".join(str(node.text or "") for node in run.xpath(".//w:t"))


def paragraph_xml_font_sizes_pt(paragraph: Any) -> list[float]:
    fallback: float | None = None
    if paragraph.style is not None:
        for style in iter_style_chain(paragraph.style):
            if style.font.size is not None:
                fallback = float(style.font.size.pt)
                break
    values: list[float] = []
    for run in paragraph._p.xpath(".//w:r"):
        if not visible_run_text(run).strip():
            continue
        size_nodes = run.xpath("./w:rPr/w:sz")
        value: float | None = fallback
        if size_nodes:
            raw = size_nodes[0].get(qn("w:val"))
            try:
                value = float(raw) / 2
            except (TypeError, ValueError):
                value = None
        if value is not None:
            values.append(value)
    return values


def paragraph_xml_font_names(paragraph: Any) -> list[str]:
    fallback: str | None = None
    if paragraph.style is not None:
        for style in iter_style_chain(paragraph.style):
            if style.font.name:
                fallback = str(style.font.name)
                break
    values: list[str] = []
    for run in paragraph._p.xpath(".//w:r"):
        if not visible_run_text(run).strip():
            continue
        font_nodes = run.xpath("./w:rPr/w:rFonts")
        if not font_nodes:
            if fallback:
                values.append(fallback)
            continue
        fonts = {
            str(font_nodes[0].get(qn(f"w:{key}")) or "").strip()
            for key in ("ascii", "hAnsi", "eastAsia", "cs")
        }
        fonts.discard("")
        values.extend(sorted(fonts))
    return values


def typography_record(
    paragraph: Any,
    *,
    location: str,
    expected_font_name: str,
    expected_font_size: float,
) -> tuple[list[dict[str, object]], dict[str, object]]:
    sizes = paragraph_xml_font_sizes_pt(paragraph)
    names = paragraph_xml_font_names(paragraph)
    record = {
        "location": location,
        "style": paragraph.style.name if paragraph.style is not None else "",
        "text_preview": paragraph.text.strip()[:100],
        "expected_font_name": expected_font_name,
        "expected_font_size_pt": expected_font_size,
        "actual_font_names": sorted(set(names)),
        "actual_font_sizes_pt": sorted(set(sizes)),
    }
    issues: list[dict[str, object]] = []
    if not sizes or any(abs(size - expected_font_size) > 0.01 for size in sizes):
        issues.append({**record, "code": "VISIBLE_FONT_SIZE_MISMATCH"})
    if not names or any(name.casefold() != expected_font_name.casefold() for name in names):
        issues.append({**record, "code": "VISIBLE_FONT_NAME_MISMATCH"})
    return issues, record


def table_border_values(table: Any) -> dict[str, str]:
    values: dict[str, str] = {}
    nodes = table._tbl.tblPr.findall(qn("w:tblBorders"))
    if not nodes:
        return values
    borders = nodes[0]
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        edge_nodes = borders.findall(qn(f"w:{edge}"))
        if edge_nodes:
            values[edge] = str(edge_nodes[0].get(qn("w:val")) or "").casefold()
    return values


def border_is_visible(value: str | None) -> bool:
    return bool(value) and value not in {"nil", "none", "0"}


def first_row_repeats(table: Any) -> bool:
    if not table.rows:
        return False
    row_properties = table.rows[0]._tr.trPr
    if row_properties is None:
        return False
    nodes = row_properties.findall(qn("w:tblHeader"))
    if not nodes:
        return False
    value = str(nodes[0].get(qn("w:val")) or "true").casefold()
    return value not in {"0", "false", "off", "none"}


def table_has_cell_shading(table: Any) -> bool:
    for shading in table._tbl.xpath(".//w:tcPr/w:shd"):
        fill = str(shading.get(qn("w:fill")) or "").strip().upper()
        if fill not in {"", "AUTO", "FFFFFF", "CLEAR", "NIL"}:
            return True
    return False


def table_header_is_bold(table: Any) -> bool:
    if not table.rows:
        return False
    visible_runs = []
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            visible_runs.extend(run for run in paragraph.runs if run.text.strip())
    return bool(visible_runs) and all(
        effective_bold(run, run._parent) for run in visible_runs
    )


def audit(
    path: Path,
    *,
    expected_line_spacing: object = "double",
    expected_font_name: str = "Times New Roman",
    expected_body_font_size: float = 12.0,
    expected_title_font_size: float = 15.0,
    expected_table_font_size: float | None = 10.0,
    expected_table_line_spacing: object = "single",
    expected_table_rule_scheme: str = "three-line",
    body_style_names: set[str] | None = None,
) -> dict[str, object]:
    if not expected_font_name.strip():
        raise ValueError("expected font name must be non-empty")
    resolved_table_font_size = (
        10.0
        if expected_table_font_size is None
        else expected_table_font_size
    )
    if min(expected_body_font_size, expected_title_font_size, resolved_table_font_size) <= 0:
        raise ValueError("expected font sizes must be positive")
    spacing_spec = parse_line_spacing_spec(expected_line_spacing)
    table_spacing_spec = parse_line_spacing_spec(expected_table_line_spacing)
    if expected_table_rule_scheme not in {
        "three-line",
        "full-grid",
        "preserve-official",
    }:
        raise ValueError("unsupported expected table rule scheme")
    body_styles = {
        *(normalize_style_token(value) for value in DEFAULT_BODY_STYLES),
        normalize_style_token("Manuscript Body"),
        *(body_style_names or set()),
    }
    document = Document(str(path))
    body = document.element.body
    paragraphs = [
        Paragraph(node, document._body)
        for node in list(body)
        if node.tag == qn("w:p")
    ]
    roles: dict[Any, str | None] = {}
    credit_block_active = False
    for paragraph in paragraphs:
        role = classify_role(paragraph, body_styles)
        if role == "credit-heading":
            credit_block_active = True
        elif role == "blank":
            pass
        elif role == "body" and credit_block_active:
            role = "credit-entry"
        elif role == "credit-entry":
            pass
        elif role == "credit-inline":
            credit_block_active = False
        else:
            credit_block_active = False
        roles[paragraph._p] = role
    issues: list[dict[str, object]] = []
    inspected: list[dict[str, object]] = []
    typography_inspected: list[dict[str, object]] = []

    for index, paragraph in enumerate(paragraphs, start=1):
        if not paragraph.text.strip():
            continue
        role = roles[paragraph._p]
        expected_size = (
            expected_title_font_size if role == "title" else expected_body_font_size
        )
        paragraph_issues, record = typography_record(
            paragraph,
            location=f"paragraph {index}",
            expected_font_name=expected_font_name,
            expected_font_size=expected_size,
        )
        issues.extend(paragraph_issues)
        typography_inspected.append(record)

    seen_cells: set[Any] = set()
    for table_index, table in enumerate(document.tables, start=1):
        table_record = {"table": table_index}
        if expected_table_rule_scheme != "preserve-official":
            borders = table_border_values(table)
            visible = {
                edge: border_is_visible(borders.get(edge))
                for edge in ("top", "bottom", "left", "right", "insideH", "insideV")
            }
            if expected_table_rule_scheme == "three-line":
                if not visible["top"] or not visible["bottom"]:
                    issues.append(
                        {
                            **table_record,
                            "code": "TABLE_OUTER_HORIZONTAL_RULES_MISSING",
                            "detail": borders,
                        }
                    )
                if any(visible[edge] for edge in ("left", "right", "insideV", "insideH")):
                    issues.append(
                        {
                            **table_record,
                            "code": "TABLE_NON_MINIMAL_RULES",
                            "detail": borders,
                        }
                    )
                header_bottom_values = []
                if table.rows:
                    for cell in table.rows[0].cells:
                        nodes = cell._tc.xpath("./w:tcPr/w:tcBorders/w:bottom")
                        header_bottom_values.append(
                            str(nodes[0].get(qn("w:val")) or "").casefold()
                            if nodes
                            else ""
                        )
                if not header_bottom_values or not all(
                    border_is_visible(value) for value in header_bottom_values
                ):
                    issues.append(
                        {
                            **table_record,
                            "code": "TABLE_HEADER_RULE_MISSING",
                            "detail": header_bottom_values,
                        }
                    )
            elif not all(visible.values()):
                issues.append(
                    {
                        **table_record,
                        "code": "TABLE_FULL_GRID_INCOMPLETE",
                        "detail": borders,
                    }
                )
            if not first_row_repeats(table):
                issues.append(
                    {**table_record, "code": "TABLE_HEADER_NOT_REPEATING"}
                )
            if not table_header_is_bold(table):
                issues.append({**table_record, "code": "TABLE_HEADER_NOT_BOLD"})
            if table_has_cell_shading(table):
                issues.append({**table_record, "code": "TABLE_CELL_SHADING"})
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    location = (
                        f"table {table_index} row {row_index} cell {cell_index} "
                        f"paragraph {paragraph_index}"
                    )
                    table_record = {
                        "location": location,
                        "style": (
                            paragraph.style.name
                            if paragraph.style is not None
                            else ""
                        ),
                        "text_preview": paragraph.text.strip()[:100],
                    }
                    actual_spacing = effective_line_spacing(paragraph)
                    if not line_spacing_matches(actual_spacing, table_spacing_spec):
                        issues.append(
                            {
                                **table_record,
                                "code": "TABLE_LINE_SPACING_MISMATCH",
                                "detail": {
                                    "expected": table_spacing_spec,
                                    "actual": actual_spacing,
                                },
                            }
                        )
                    for field, label in (
                        ("space_before", "BEFORE"),
                        ("space_after", "AFTER"),
                    ):
                        automatic = automatic_spacing_sources(paragraph, field)
                        if automatic:
                            issues.append(
                                {
                                    **table_record,
                                    "code": f"TABLE_AUTOSPACING_{label}",
                                    "detail": automatic,
                                }
                            )
                        points, source = effective_spacing(paragraph, field)
                        if abs(points) > 0.01:
                            issues.append(
                                {
                                    **table_record,
                                    "code": f"TABLE_SPACE_{label}",
                                    "detail": {
                                        "points": round(points, 3),
                                        "source": source,
                                    },
                                }
                            )
                    if paragraph.text.strip():
                        paragraph_issues, record = typography_record(
                            paragraph,
                            location=location,
                            expected_font_name=expected_font_name,
                            expected_font_size=resolved_table_font_size,
                        )
                        issues.extend(paragraph_issues)
                        typography_inspected.append(record)

    same_size_roles = {
        "authors",
        "affiliation",
        "author_note",
        "correspondence",
        "orcid",
        "keywords",
        "heading",
        "body",
        "declaration-inline",
        "credit-heading",
        "credit-inline",
        "credit-entry",
        "list-item",
    }
    for index, paragraph in enumerate(paragraphs, start=1):
        role = roles[paragraph._p]
        if role is None:
            continue
        record = {
            "paragraph": index,
            "role": role,
            "style": paragraph.style.name if paragraph.style is not None else "",
            "text_preview": paragraph.text.strip()[:100],
        }
        inspected.append(record)

        actual_spacing = effective_line_spacing(paragraph)
        if not line_spacing_matches(actual_spacing, spacing_spec):
            issues.append(
                {
                    **record,
                    "code": "GLOBAL_LINE_SPACING_MISMATCH",
                    "detail": {"expected": spacing_spec, "actual": actual_spacing},
                }
            )
        for field, label in (("space_before", "BEFORE"), ("space_after", "AFTER")):
            automatic = automatic_spacing_sources(paragraph, field)
            if automatic:
                issues.append(
                    {
                        **record,
                        "code": f"SEMANTIC_AUTOSPACING_{label}",
                        "detail": automatic,
                    }
                )
            points, source = effective_spacing(paragraph, field)
            if abs(points) > 0.01:
                issues.append(
                    {
                        **record,
                        "code": f"SEMANTIC_SPACE_{label}",
                        "detail": {"points": round(points, 3), "source": source},
                    }
                )

        if role in same_size_roles and paragraph.text.strip():
            sizes = paragraph_font_sizes_pt(paragraph)
            if not sizes or any(
                abs(size - expected_body_font_size) > 0.01 for size in sizes
            ):
                issues.append(
                    {
                        **record,
                        "code": "BODY_SIZE_ROLE_MISMATCH",
                        "detail": {
                            "expected_pt": expected_body_font_size,
                            "actual_pt": sizes,
                        },
                    }
                )

        if role == "keywords":
            for issue in label_weight_issues(
                paragraph, KEYWORD_LABEL_RE, prefix="KEYWORDS"
            ):
                issues.append({**record, **issue})
            before = adjacent_blank_count(document, paragraph._p, -1)
            after = adjacent_blank_count(document, paragraph._p, 1)
            if before != 0:
                issues.append(
                    {**record, "code": "KEYWORDS_BLANK_BEFORE", "detail": before}
                )
            if after != 1:
                issues.append(
                    {**record, "code": "KEYWORDS_BLANK_AFTER", "detail": after}
                )

        if role in {"declaration-inline", "credit-inline"}:
            pattern = (
                CREDIT_INLINE_RE
                if role == "credit-inline"
                else DECLARATION_INLINE_RE
            )
            prefix = "CREDIT" if role == "credit-inline" else "DECLARATION"
            for issue in label_weight_issues(
                paragraph, pattern, prefix=prefix
            ):
                issues.append({**record, **issue})
            if role == "credit-inline" and not credit_role_labels(paragraph.text):
                issues.append({**record, "code": "CREDIT_ROLE_VOCABULARY_MISSING"})

        if role in {"heading", "credit-heading"}:
            after = adjacent_blank_count(document, paragraph._p, 1)
            if after != 0:
                issues.append(
                    {**record, "code": "HEADING_BLANK_AFTER", "detail": after}
                )
            if re.fullmatch(r"abstract", paragraph.text.strip(), re.I):
                continue

        if role in {
            "heading",
            "credit-heading",
            "declaration-inline",
            "credit-inline",
        }:
            previous_role = previous_nonblank_role(document, paragraph._p, roles)
            expected_before = (
                0 if previous_role in {"heading", "credit-heading"} else 1
            )
            actual_before = adjacent_blank_count(document, paragraph._p, -1)
            if actual_before != expected_before:
                issues.append(
                    {
                        **record,
                        "code": "SEMANTIC_BLANK_BEFORE",
                        "detail": {
                            "expected": expected_before,
                            "actual": actual_before,
                            "previous_role": previous_role,
                        },
                    }
                )

        if role == "credit-heading":
            if next_nonblank_role(document, paragraph._p, roles) != "credit-entry":
                issues.append({**record, "code": "CREDIT_STATEMENT_MISSING"})

        if role == "credit-entry":
            if not credit_role_labels(paragraph.text):
                issues.append(
                    {**record, "code": "CREDIT_ENTRY_WITHOUT_STANDARD_ROLE"}
                )
            if next_nonblank_role(document, paragraph._p, roles) == "credit-entry":
                between = adjacent_blank_count(document, paragraph._p, 1)
                if between != 0:
                    issues.append(
                        {
                            **record,
                            "code": "CREDIT_ENTRY_BLANK_BETWEEN",
                            "detail": between,
                        }
                    )

        if role == "list-item":
            previous_role = previous_nonblank_role(document, paragraph._p, roles)
            next_role = next_nonblank_role(document, paragraph._p, roles)
            before = adjacent_blank_count(document, paragraph._p, -1)
            after = adjacent_blank_count(document, paragraph._p, 1)
            if previous_role == "list-item" and before != 0:
                issues.append(
                    {**record, "code": "LIST_ITEM_BLANK_BETWEEN", "detail": before}
                )
            elif previous_role == "heading" and before != 0:
                issues.append(
                    {**record, "code": "LIST_BLOCK_BLANK_AFTER_HEADING", "detail": before}
                )
            elif previous_role not in {None, "heading", "list-item"} and before != 1:
                issues.append(
                    {
                        **record,
                        "code": "LIST_BLOCK_BLANK_BEFORE",
                        "detail": {"expected": 1, "actual": before},
                    }
                )
            if next_role == "list-item" and after != 0:
                issues.append(
                    {**record, "code": "LIST_ITEM_BLANK_BETWEEN", "detail": after}
                )
            elif next_role not in {None, "heading", "list-item"} and after != 1:
                issues.append(
                    {
                        **record,
                        "code": "LIST_BLOCK_BLANK_AFTER",
                        "detail": {"expected": 1, "actual": after},
                    }
                )

    return {
        "status": "SEMANTIC_RHYTHM_PASS" if not issues else "FAIL",
        "document": str(path.resolve()),
        "expected_line_spacing": spacing_spec["label"],
        "expected_font_name": expected_font_name,
        "expected_body_font_size_pt": expected_body_font_size,
        "expected_title_font_size_pt": expected_title_font_size,
        "expected_table_font_size_pt": resolved_table_font_size,
        "expected_table_line_spacing": table_spacing_spec["label"],
        "expected_table_rule_scheme": expected_table_rule_scheme,
        "issue_count": len(issues),
        "issues": issues,
        "inspected": inspected,
        "typography_inspected": typography_inspected,
        "boundary": (
            "Checks manuscript typography and semantic vertical rhythm only; "
            "rendered-page and journal-source gates remain independent."
        ),
    }


def render_text(report: dict[str, object]) -> str:
    lines = [
        f"DOCX semantic-rhythm audit: {report['status']}",
        f"Document: {report['document']}",
        f"Issues: {report['issue_count']}",
    ]
    for issue in report["issues"]:
        lines.append(
            f"ERROR: paragraph {issue.get('paragraph', '?')}: "
            f"{issue['code']} - {issue.get('detail', '')}"
        )
    lines.append(f"BOUNDARY: {report['boundary']}")
    return "\n".join(lines)


def main() -> int:
    args = parse_args()
    try:
        report = audit(
            args.document,
            expected_line_spacing=args.expected_line_spacing,
            expected_font_name=args.expected_font_name,
            expected_body_font_size=args.expected_body_font_size,
            expected_title_font_size=args.expected_title_font_size,
            expected_table_font_size=args.expected_table_font_size,
            expected_table_line_spacing=args.expected_table_line_spacing,
            expected_table_rule_scheme=args.expected_table_rule_scheme,
            body_style_names={normalize_style_token(value) for value in args.body_style},
        )
    except (OSError, ValueError) as exc:
        print(f"Unable to audit DOCX semantic rhythm: {exc}", file=sys.stderr)
        return 2
    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(
            json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
        )
    if args.json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    else:
        print(render_text(report))
    return 0 if report["status"] == "SEMANTIC_RHYTHM_PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
