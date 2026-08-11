#!/usr/bin/env python3
"""Apply a restrained, journal-neutral manuscript DOCX profile."""

from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - environment-dependent error path
    print(
        "python-docx is required. Use the bundled workspace Python runtime.",
        file=sys.stderr,
    )
    raise SystemExit(2) from exc

from audit_docx_front_matter import (  # noqa: E402
    ALIGNMENT_VALUES,
    ROLE_ORDER,
    build_explicit_paragraph_map,
    build_explicit_style_map,
    classify_role,
)
from audit_docx_manuscript_style import (  # noqa: E402
    DEFAULT_BODY_STYLES,
    normalize_style_token,
    paragraph_is_body,
    paragraph_is_nonbody,
    paragraph_is_structurally_empty,
    parse_line_spacing_spec,
)


from docx_semantic_rules import (  # noqa: E402
    BODY_FONT_SIZE_PT,
    CREDIT_HEADING_RE,
    DECLARATION_HEADING_RE,
    DECLARATION_INLINE_RE,
    HEADING_STYLE_RE,
    KEYWORD_LABEL_RE,
    bold_leading_label,
)


TITLE_FONT_SIZE_PT = 15.0
TABLE_FONT_SIZE_PT = 10.0
DEFAULT_FONT_NAME = "Times New Roman"


def profile_styles(
    body_font_size: float, title_font_size: float
) -> dict[str, tuple[str, float, bool, float]]:
    return {
        "title": ("Manuscript Title", title_font_size, True, 1.0),
        "authors": ("Manuscript Authors", body_font_size, False, 1.0),
        "affiliation": ("Manuscript Affiliation", body_font_size, False, 1.0),
        "author_note": ("Manuscript Author Note", body_font_size, False, 1.0),
        "correspondence": (
            "Manuscript Correspondence",
            body_font_size,
            False,
            1.0,
        ),
        "orcid": ("Manuscript ORCID", body_font_size, False, 1.0),
        "keywords": ("Manuscript Keywords", body_font_size, False, 1.0),
        "heading": ("Manuscript Heading", body_font_size, True, 1.0),
        "credit-entry": ("Manuscript CRediT Entry", body_font_size, False, 1.0),
        "body": ("Manuscript Body", body_font_size, False, 2.0),
        "reference": ("Manuscript Reference", body_font_size, False, 1.0),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Normalize a scientific DOCX to a restrained journal-neutral profile "
            "without changing text content."
        )
    )
    parser.add_argument("document", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    parser.add_argument(
        "--front-matter-alignment",
        choices=tuple(ALIGNMENT_VALUES),
        default="left",
    )
    parser.add_argument("--line-spacing", default="double")
    parser.add_argument("--font-name", default=DEFAULT_FONT_NAME)
    parser.add_argument(
        "--body-font-size", type=float, default=BODY_FONT_SIZE_PT
    )
    parser.add_argument(
        "--title-font-size", type=float, default=TITLE_FONT_SIZE_PT
    )
    parser.add_argument(
        "--table-font-size",
        type=float,
        default=TABLE_FONT_SIZE_PT,
        help="Table-cell font size; defaults to 10 pt.",
    )
    parser.add_argument(
        "--table-line-spacing",
        default="single",
        help="Table-cell line spacing; defaults to single.",
    )
    parser.add_argument(
        "--abstract-start", choices=("integrated", "new-page"), default="integrated"
    )
    parser.add_argument(
        "--body-style",
        action="append",
        default=[],
        help="Existing body-prose style name or ID. Repeat as needed.",
    )
    for role in ROLE_ORDER:
        option_role = role.replace("_", "-")
        style_flags = [f"--{option_role}-style"]
        paragraph_flags = [f"--{option_role}-paragraph"]
        if option_role != role:
            style_flags.append(f"--{role}-style")
            paragraph_flags.append(f"--{role}-paragraph")
        parser.add_argument(
            *style_flags,
            dest=f"{role}_style",
            action="append",
            default=[],
            help=f"Existing paragraph style name or ID for {role}.",
        )
        parser.add_argument(
            *paragraph_flags,
            dest=f"{role}_paragraph",
            action="append",
            type=int,
            default=[],
            help=(
                f"One-based top-level paragraph number for {role}. "
                "Repeat for multi-paragraph roles."
            ),
        )
    return parser.parse_args()


def text_node_values(document: Any) -> list[str]:
    return [str(node.text or "") for node in document.element.body.xpath(".//w:t")]


def set_font_name(font: Any, name: str) -> None:
    font.name = name
    r_pr = font._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), name)


def set_line_spacing(paragraph_format: Any, spec: dict[str, object]) -> None:
    kind = str(spec["kind"])
    value = float(spec["value"])
    if kind == "multiple":
        paragraph_format.line_spacing = value
        paragraph_format.line_spacing_rule = None
    elif kind == "exact":
        paragraph_format.line_spacing = Pt(value)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    elif kind == "at-least":
        paragraph_format.line_spacing = Pt(value)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    else:  # pragma: no cover - guarded by parse_line_spacing_spec
        raise ValueError(f"Unsupported line-spacing kind: {kind}")


def remove_auto_spacing(paragraph_properties: Any) -> None:
    if paragraph_properties is None:
        return
    for spacing in paragraph_properties.xpath("./w:spacing"):
        spacing.attrib.pop(qn("w:beforeAutospacing"), None)
        spacing.attrib.pop(qn("w:afterAutospacing"), None)


def ensure_style(
    document: Any,
    name: str,
    *,
    size_pt: float,
    bold: bool,
    line_spacing: dict[str, object],
    alignment: Any,
    font_name: str = DEFAULT_FONT_NAME,
) -> Any:
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = None
    set_font_name(style.font, font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.alignment = alignment
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    set_line_spacing(style.paragraph_format, line_spacing)
    remove_auto_spacing(style.element.get_or_add_pPr())
    return style


def format_paragraph(
    paragraph: Any,
    style: Any,
    *,
    size_pt: float,
    bold: bool,
    alignment: Any,
    line_spacing: dict[str, object],
) -> None:
    paragraph.style = style
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_line_spacing(paragraph.paragraph_format, line_spacing)
    remove_auto_spacing(paragraph._p.get_or_add_pPr())
    font_name = style.font.name or DEFAULT_FONT_NAME
    for run in paragraph.runs:
        set_font_name(run.font, font_name)
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)


def format_run_xml_typography(
    run: Any, *, font_name: str, font_size: float
) -> None:
    r_pr = run.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run.insert(0, r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), font_name)
    half_points = str(int(round(font_size * 2)))
    for tag in ("w:sz", "w:szCs"):
        nodes = r_pr.findall(qn(tag))
        node = nodes[0] if nodes else OxmlElement(tag)
        if not nodes:
            r_pr.append(node)
        node.set(qn("w:val"), half_points)
        for duplicate in nodes[1:]:
            r_pr.remove(duplicate)
    color_nodes = r_pr.findall(qn("w:color"))
    color = color_nodes[0] if color_nodes else OxmlElement("w:color")
    if not color_nodes:
        r_pr.append(color)
    color.set(qn("w:val"), "000000")
    for duplicate in color_nodes[1:]:
        r_pr.remove(duplicate)


def apply_visible_typography(
    document: Any,
    *,
    title_nodes: set[Any],
    font_name: str,
    body_font_size: float,
    title_font_size: float,
    table_font_size: float,
    line_spacing: dict[str, object],
    table_line_spacing: dict[str, object],
) -> dict[str, int]:
    """Apply the resolved font contract to every visible manuscript run."""

    top_level_count = 0
    for paragraph in document.paragraphs:
        size = title_font_size if paragraph._p in title_nodes else body_font_size
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        set_line_spacing(paragraph.paragraph_format, line_spacing)
        remove_auto_spacing(paragraph._p.get_or_add_pPr())
        if paragraph.text.strip():
            top_level_count += 1
        for run in paragraph._p.xpath(".//w:r"):
            format_run_xml_typography(
                run, font_name=font_name, font_size=size
            )

    table_count = 0
    seen_cells: set[int] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    set_line_spacing(
                        paragraph.paragraph_format, table_line_spacing
                    )
                    remove_auto_spacing(paragraph._p.get_or_add_pPr())
                    if paragraph.text.strip():
                        table_count += 1
                    for run in paragraph._p.xpath(".//w:r"):
                        format_run_xml_typography(
                            run,
                            font_name=font_name,
                            font_size=table_font_size,
                        )
    return {
        "top_level_visible_paragraphs": top_level_count,
        "table_visible_paragraphs": table_count,
    }


def paragraph_is_abstract(paragraph: Any) -> bool:
    return bool(re.fullmatch(r"abstract", paragraph.text.strip(), re.I))


def flatten_recognized_front_matter_tables(
    document: Any, explicit_styles: dict[str, set[str]]
) -> int:
    body = document.element.body
    flattened = 0
    first_nonempty = True
    for child in list(body):
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document._body)
            if paragraph_is_abstract(paragraph):
                break
            if paragraph.text.strip():
                first_nonempty = False
            if child.xpath(".//w:txbxContent"):
                raise ValueError(
                    "Front matter in a text box cannot be safely normalized. "
                    "Convert it to ordinary paragraphs first."
                )
            continue
        if child.tag != qn("w:tbl"):
            continue

        paragraphs = [
            Paragraph(node, document._body)
            for node in child.xpath(".//w:tr/w:tc/w:p")
            if Paragraph(node, document._body).text.strip()
        ]
        if not paragraphs:
            body.remove(child)
            flattened += 1
            continue
        roles: list[str | None] = []
        for index, paragraph in enumerate(paragraphs):
            roles.append(
                classify_role(
                    paragraph,
                    explicit_styles,
                    first_nonempty=first_nonempty and index == 0,
                )
            )
        if any(role is None for role in roles):
            raise ValueError(
                "A table before Abstract contains non-front-matter content; "
                "refusing to flatten it automatically."
            )
        for paragraph in paragraphs:
            body.insert(body.index(child), deepcopy(paragraph._p))
        body.remove(child)
        flattened += 1
        first_nonempty = False
    return flattened


def front_matter_records(
    document: Any,
    explicit_styles: dict[str, set[str]],
    explicit_paragraphs: dict[str, set[int]],
) -> tuple[list[tuple[Any, str]], Any | None]:
    records: list[tuple[Any, str]] = []
    abstract: Any | None = None
    first_nonempty = True
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph_is_abstract(paragraph):
            abstract = paragraph
            break
        if not paragraph.text.strip():
            continue
        role = classify_role(
            paragraph,
            explicit_styles,
            first_nonempty=first_nonempty,
            paragraph_number=paragraph_number,
            explicit_paragraphs=explicit_paragraphs,
        )
        first_nonempty = False
        if role is not None:
            records.append((paragraph, role))
    return records, abstract


def normalize_front_matter_blanks(
    document: Any,
    records: list[tuple[Any, str]],
    abstract: Any | None,
    abstract_start: str,
    separator_style: Any,
    separator_spacing: dict[str, object],
) -> None:
    if not records or abstract is None:
        return
    body = document.element.body
    record_nodes = [paragraph._p for paragraph, _ in records]
    first_index = min(body.index(node) for node in record_nodes)
    last_index = max(body.index(node) for node in record_nodes)
    abstract_index = body.index(abstract._p)

    for child in list(body)[first_index + 1 : last_index]:
        if child.tag == qn("w:p") and not Paragraph(child, document._body).text.strip():
            body.remove(child)

    # Insert exactly one real Enter-created paragraph between adjacent present
    # semantic blocks. Consecutive paragraphs with the same role stay compact.
    ordered_records = sorted(records, key=lambda item: body.index(item[0]._p))
    previous_role = ordered_records[0][1]
    for paragraph, role in ordered_records[1:]:
        if role != previous_role:
            blank = OxmlElement("w:p")
            body.insert(body.index(paragraph._p), blank)
            format_paragraph(
                Paragraph(blank, document._body),
                separator_style,
                size_pt=BODY_FONT_SIZE_PT,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=separator_spacing,
            )
        previous_role = role

    abstract_index = body.index(abstract._p)
    last_index = max(body.index(node) for node in record_nodes)
    for child in list(body)[last_index + 1 : abstract_index]:
        if child.tag == qn("w:p") and not Paragraph(child, document._body).text.strip():
            body.remove(child)

    blank = OxmlElement("w:p")
    body.insert(body.index(abstract._p), blank)
    format_paragraph(
        Paragraph(blank, document._body),
        separator_style,
        size_pt=BODY_FONT_SIZE_PT,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=separator_spacing,
    )
    if abstract_start == "integrated":
        abstract.paragraph_format.page_break_before = False
    else:
        abstract.paragraph_format.page_break_before = True


def normalize_body_separators(
    document: Any,
    body_paragraphs: list[Any],
    body_style: Any,
    line_spacing: dict[str, object],
) -> None:
    body = document.element.body
    body_nodes = {id(paragraph._p): paragraph._p for paragraph in body_paragraphs}
    children = list(body)
    index = 0
    while index < len(children):
        current = children[index]
        if id(current) not in body_nodes:
            index += 1
            continue
        cursor = index + 1
        blanks: list[Any] = []
        while cursor < len(children):
            candidate = children[cursor]
            if candidate.tag != qn("w:p"):
                break
            paragraph = Paragraph(candidate, document._body)
            if paragraph_is_structurally_empty(paragraph):
                blanks.append(candidate)
                cursor += 1
                continue
            break
        if cursor < len(children) and id(children[cursor]) in body_nodes:
            if not blanks:
                blank = OxmlElement("w:p")
                body.insert(body.index(children[cursor]), blank)
                blanks = [blank]
            for duplicate in blanks[1:]:
                body.remove(duplicate)
            separator = Paragraph(blanks[0], document._body)
            format_paragraph(
                separator,
                body_style,
                size_pt=12.0,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=line_spacing,
            )
        children = list(body)
        index += 1


def adjacent_blank_nodes_before(document: Any, node: Any) -> list[Any]:
    body = document.element.body
    children = list(body)
    index = children.index(node) - 1
    blanks: list[Any] = []
    while index >= 0:
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            break
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            break
        blanks.insert(0, candidate)
        index -= 1
    return blanks


def adjacent_blank_nodes_after(document: Any, node: Any) -> list[Any]:
    body = document.element.body
    children = list(body)
    index = children.index(node) + 1
    blanks: list[Any] = []
    while index < len(children):
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            break
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            break
        blanks.append(candidate)
        index += 1
    return blanks


def format_separator(node: Any, document: Any, style: Any, line_spacing: dict[str, object]) -> None:
    format_paragraph(
        Paragraph(node, document._body),
        style,
        size_pt=BODY_FONT_SIZE_PT,
        bold=False,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=line_spacing,
    )


def set_blank_count_before(
    document: Any,
    node: Any,
    desired: int,
    separator_style: Any,
    line_spacing: dict[str, object],
) -> None:
    body = document.element.body
    blanks = adjacent_blank_nodes_before(document, node)
    for duplicate in blanks[desired:]:
        body.remove(duplicate)
    blanks = blanks[:desired]
    while len(blanks) < desired:
        blank = OxmlElement("w:p")
        body.insert(body.index(node), blank)
        blanks.append(blank)
    for blank in blanks:
        format_separator(blank, document, separator_style, line_spacing)


def set_blank_count_after(
    document: Any,
    node: Any,
    desired: int,
    separator_style: Any,
    line_spacing: dict[str, object],
) -> None:
    body = document.element.body
    blanks = adjacent_blank_nodes_after(document, node)
    for duplicate in blanks[desired:]:
        body.remove(duplicate)
    blanks = blanks[:desired]
    while len(blanks) < desired:
        blank = OxmlElement("w:p")
        body.insert(body.index(node) + 1, blank)
        blanks.insert(0, blank)
    for blank in blanks:
        format_separator(blank, document, separator_style, line_spacing)


def previous_nonblank_paragraph(document: Any, node: Any) -> Any | None:
    body = document.element.body
    children = list(body)
    index = children.index(node) - 1
    while index >= 0:
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            return None
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            return candidate
        index -= 1
    return None


def next_nonblank_paragraph(document: Any, node: Any) -> Any | None:
    body = document.element.body
    children = list(body)
    index = children.index(node) + 1
    while index < len(children):
        candidate = children[index]
        if candidate.tag != qn("w:p"):
            return None
        paragraph = Paragraph(candidate, document._body)
        if not paragraph_is_structurally_empty(paragraph):
            return candidate
        index += 1
    return None


def normalize_semantic_vertical_rhythm(
    document: Any,
    *,
    heading_nodes: set[Any],
    abstract_node: Any | None,
    keyword_nodes: set[Any],
    declaration_inline_nodes: set[Any],
    credit_entry_nodes: set[Any],
    separator_style: Any,
    line_spacing: dict[str, object],
) -> None:
    """Enforce the manuscript's semantic blank-line matrix."""
    body = document.element.body
    ordered = [node for node in list(body) if node.tag == qn("w:p")]

    # Headings use the global line spacing and sit directly on their first paragraph.
    for node in ordered:
        if node in heading_nodes:
            set_blank_count_after(
                document, node, 0, separator_style, line_spacing
            )

    # Keywords follow the abstract directly and are followed by one natural empty line.
    for node in ordered:
        if node in keyword_nodes:
            set_blank_count_before(
                document, node, 0, separator_style, line_spacing
            )
            set_blank_count_after(
                document, node, 1, separator_style, line_spacing
            )

    # A CRediT block is compact. Preserve supplied author-entry paragraphs, but
    # remove empty paragraphs between consecutive entries rather than treating
    # them as ordinary body prose.
    for node in ordered:
        if node not in credit_entry_nodes:
            continue
        next_node = next_nonblank_paragraph(document, node)
        if next_node in credit_entry_nodes:
            set_blank_count_after(
                document, node, 0, separator_style, line_spacing
            )

    # Every new section/subsection/declaration block has exactly one empty line
    # before it, except Abstract (handled by the front-matter contract) and
    # consecutive headings, which remain together.
    for node in ordered:
        if node == abstract_node:
            continue
        if node not in heading_nodes and node not in declaration_inline_nodes:
            continue
        previous = previous_nonblank_paragraph(document, node)
        desired = 0 if previous in heading_nodes else 1
        set_blank_count_before(
            document, node, desired, separator_style, line_spacing
        )


def apply_profile(
    document_path: Path,
    output_path: Path,
    *,
    front_matter_alignment: str,
    line_spacing_token: str,
    font_name: str,
    body_font_size: float,
    title_font_size: float,
    table_font_size: float | None,
    table_line_spacing_token: str,
    abstract_start: str,
    body_style_tokens: set[str],
    explicit_styles: dict[str, set[str]],
    explicit_paragraphs: dict[str, set[int]],
) -> dict[str, object]:
    if document_path.resolve() == output_path.resolve():
        raise ValueError("--out must differ from the input path; preserve the source DOCX.")
    if not font_name.strip():
        raise ValueError("font name must be non-empty")
    if body_font_size <= 0 or title_font_size <= 0:
        raise ValueError("body and title font sizes must be positive")
    resolved_table_font_size = (
        TABLE_FONT_SIZE_PT if table_font_size is None else table_font_size
    )
    if resolved_table_font_size <= 0:
        raise ValueError("table font size must be positive")
    document = Document(str(document_path))
    original_text_stream = "".join(text_node_values(document))
    flattened_tables = flatten_recognized_front_matter_tables(
        document, explicit_styles
    )

    alignment = ALIGNMENT_VALUES[front_matter_alignment]
    body_spacing = parse_line_spacing_spec(line_spacing_token)
    table_spacing = parse_line_spacing_spec(table_line_spacing_token)
    profile_definitions = profile_styles(body_font_size, title_font_size)
    styles: dict[str, Any] = {}
    for role, (name, size, bold, _) in profile_definitions.items():
        role_alignment = alignment if role in ROLE_ORDER else WD_ALIGN_PARAGRAPH.LEFT
        styles[role] = ensure_style(
            document,
            name,
            size_pt=size,
            bold=bold,
            line_spacing=body_spacing,
            alignment=role_alignment,
            font_name=font_name,
        )

    records, abstract = front_matter_records(
        document, explicit_styles, explicit_paragraphs
    )
    role_counts = {role: 0 for role in ROLE_ORDER}
    for paragraph, role in records:
        role_counts[role] += 1
        _, size, bold, _ = profile_definitions[role]
        format_paragraph(
            paragraph,
            styles[role],
            size_pt=size,
            bold=bold,
            alignment=alignment,
            line_spacing=body_spacing,
        )

    if abstract is not None:
        format_paragraph(
            abstract,
            styles["heading"],
            size_pt=body_font_size,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=body_spacing,
        )

    normalized_body_tokens = set(DEFAULT_BODY_STYLES)
    normalized_body_tokens.update(body_style_tokens)
    normalized_body_tokens.add(normalize_style_token("Manuscript Body"))
    excluded = {
        normalize_style_token(style.name)
        for role, style in styles.items()
        if role != "body"
    }
    body_paragraphs: list[Any] = []
    heading_nodes: set[Any] = set()
    keyword_nodes: set[Any] = set()
    declaration_inline_nodes: set[Any] = set()
    credit_entry_nodes: set[Any] = set()
    if abstract is not None:
        heading_nodes.add(abstract._p)
    abstract_seen = False
    reference_section_seen = False
    credit_block_active = False
    for paragraph in document.paragraphs:
        if paragraph_is_abstract(paragraph):
            abstract_seen = True
            continue
        if not abstract_seen:
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if re.search(r"keyword", style_name, re.I) or re.match(
            r"^\s*keywords?\s*:", paragraph.text, re.I
        ):
            format_paragraph(
                paragraph,
                styles["keywords"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            bold_leading_label(paragraph, KEYWORD_LABEL_RE)
            keyword_nodes.add(paragraph._p)
            credit_block_active = False
            continue
        if (
            HEADING_STYLE_RE.match(style_name)
            or DECLARATION_HEADING_RE.fullmatch(paragraph.text)
        ) or (
            paragraph.text.strip() and style_name.lower().startswith("heading")
        ):
            format_paragraph(
                paragraph,
                styles["heading"],
                size_pt=body_font_size,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            heading_nodes.add(paragraph._p)
            credit_block_active = bool(CREDIT_HEADING_RE.fullmatch(paragraph.text))
            if re.fullmatch(
                r"(?:references|bibliography|literature cited)",
                paragraph.text.strip(),
                re.I,
            ):
                reference_section_seen = True
            elif reference_section_seen:
                reference_section_seen = False
            continue
        if DECLARATION_INLINE_RE.match(paragraph.text):
            reference_section_seen = False
            credit_block_active = False
            format_paragraph(
                paragraph,
                styles["body"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            bold_leading_label(paragraph, DECLARATION_INLINE_RE)
            body_paragraphs.append(paragraph)
            declaration_inline_nodes.add(paragraph._p)
            continue
        paragraph_style_tokens = {
            normalize_style_token(style_name),
            normalize_style_token(
                paragraph.style.style_id if paragraph.style is not None else ""
            ),
        }
        if credit_block_active and paragraph.text.strip() and (
            paragraph_is_body(paragraph, normalized_body_tokens, excluded)
            or normalize_style_token("Manuscript CRediT Entry")
            in paragraph_style_tokens
        ):
            format_paragraph(
                paragraph,
                styles["credit-entry"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            credit_entry_nodes.add(paragraph._p)
            continue
        if reference_section_seen and paragraph.text.strip():
            credit_block_active = False
            format_paragraph(
                paragraph,
                styles["reference"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.5)
            continue
        if paragraph_is_body(paragraph, normalized_body_tokens, excluded):
            format_paragraph(
                paragraph,
                styles["body"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            body_paragraphs.append(paragraph)
            continue
        if paragraph.text.strip() and not paragraph_is_nonbody(paragraph, excluded):
            format_paragraph(
                paragraph,
                styles["body"],
                size_pt=body_font_size,
                bold=False,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                line_spacing=body_spacing,
            )
            body_paragraphs.append(paragraph)

    normalize_front_matter_blanks(
        document,
        records,
        abstract,
        abstract_start,
        styles["body"],
        body_spacing,
    )
    normalize_body_separators(
        document, body_paragraphs, styles["body"], body_spacing
    )
    normalize_semantic_vertical_rhythm(
        document,
        heading_nodes=heading_nodes,
        abstract_node=abstract._p if abstract is not None else None,
        keyword_nodes=keyword_nodes,
        declaration_inline_nodes=declaration_inline_nodes,
        credit_entry_nodes=credit_entry_nodes,
        separator_style=styles["body"],
        line_spacing=body_spacing,
    )

    title_nodes = {
        paragraph._p for paragraph, role in records if role == "title"
    }
    typography_counts = apply_visible_typography(
        document,
        title_nodes=title_nodes,
        font_name=font_name.strip(),
        body_font_size=body_font_size,
        title_font_size=title_font_size,
        table_font_size=resolved_table_font_size,
        line_spacing=body_spacing,
        table_line_spacing=table_spacing,
    )

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        vertical_nodes = section._sectPr.xpath("./w:vAlign")
        vertical = vertical_nodes[0] if vertical_nodes else OxmlElement("w:vAlign")
        if not vertical_nodes:
            section._sectPr.append(vertical)
        vertical.set(qn("w:val"), "top")

    if "".join(text_node_values(document)) != original_text_stream:
        raise ValueError("Text-node preservation check failed; output was not written.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return {
        "output": str(output_path.resolve()),
        "front_matter_alignment": front_matter_alignment,
        "line_spacing": body_spacing["label"],
        "font_name": font_name.strip(),
        "body_font_size_pt": body_font_size,
        "title_font_size_pt": title_font_size,
        "table_font_size_pt": resolved_table_font_size,
        "table_line_spacing": table_spacing["label"],
        "role_counts": role_counts,
        "flattened_front_matter_tables": flattened_tables,
        "typography_counts": typography_counts,
    }


def main() -> int:
    args = parse_args()
    try:
        explicit_styles = build_explicit_style_map(args)
        explicit_paragraphs = build_explicit_paragraph_map(args)
        body_style_tokens = {
            normalize_style_token(value) for value in args.body_style
        }
        result = apply_profile(
            args.document,
            args.out,
            front_matter_alignment=args.front_matter_alignment,
            line_spacing_token=args.line_spacing,
            font_name=args.font_name,
            body_font_size=args.body_font_size,
            title_font_size=args.title_font_size,
            table_font_size=args.table_font_size,
            table_line_spacing_token=args.table_line_spacing,
            abstract_start=args.abstract_start,
            body_style_tokens=body_style_tokens,
            explicit_styles=explicit_styles,
            explicit_paragraphs=explicit_paragraphs,
        )
    except (OSError, ValueError) as exc:
        print(f"Unable to apply manuscript profile: {exc}", file=sys.stderr)
        return 2
    print(f"Manuscript profile applied: {result}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
