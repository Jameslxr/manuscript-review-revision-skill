#!/usr/bin/env python3
"""Audit restrained styling, paragraph structure, and DOCX numbering."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - environment-dependent error path
    print(
        "python-docx is required. Use the bundled workspace Python runtime.",
        file=sys.stderr,
    )
    raise SystemExit(2) from exc


HEADING_STYLE_RE = re.compile(
    r"^(?:title|subtitle|heading\s*[1-9]|manuscript\s+(?:title|heading))$",
    re.IGNORECASE,
)
SAFE_THEME_TOKENS = {"DARK", "TEXT_1", "TEXT_2", "NONE"}
LINE_SPACING_ALIASES = {
    "single": ("multiple", 1.0),
    "1.0": ("multiple", 1.0),
    "1.15": ("multiple", 1.15),
    "1.5": ("multiple", 1.5),
    "double": ("multiple", 2.0),
    "2.0": ("multiple", 2.0),
}
PAGE_FIELD_RE = re.compile(r"(?:^|\s)PAGE(?:\s|\\|$)", re.IGNORECASE)
DEFAULT_BODY_STYLES = {
    "body",
    "body text",
    "bodytext",
    "normal",
    "text body",
    "正文",
    "正文文本",
}
NON_BODY_STYLE_RE = re.compile(
    r"(?:title|subtitle|heading|author|affiliation|correspond|keyword|declaration|"
    r"credit|contribut|caption|bibliograph|reference|"
    r"list|quote|footnote|endnote|header|footer)",
    re.IGNORECASE,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Audit explicit or theme-colored title/headings, paragraph borders, "
            "shading, body-paragraph separation, continuous line numbering, and "
            "dynamic page numbering in a DOCX submission manuscript."
        )
    )
    parser.add_argument("document", type=Path)
    parser.add_argument("--json", action="store_true", help="Emit JSON.")
    parser.add_argument(
        "--output-json", type=Path, help="Also write the report to this JSON file."
    )
    parser.add_argument(
        "--paragraph-separation",
        choices=("literal-blank",),
        default="literal-blank",
        help=(
            "Require zero body paragraph spacing plus one real empty paragraph "
            "between adjacent body paragraphs. This is a non-bypassable output "
            "invariant for this Skill."
        ),
    )
    parser.add_argument(
        "--body-style",
        action="append",
        default=[],
        help=(
            "Additional body paragraph style name or style ID. Repeat for multiple "
            "custom styles."
        ),
    )
    parser.add_argument(
        "--exclude-style",
        action="append",
        default=[],
        help=(
            "Explicit non-body paragraph style name or style ID. Repeat for "
            "custom title, affiliation, caption, bibliography, or list styles."
        ),
    )
    parser.add_argument(
        "--expected-line-spacing",
        default="double",
        help=(
            "Expected body and literal-separator line spacing. Use single, 1.15, "
            "1.5, double, a positive numeric multiple, exact:<pt>pt, "
            "or at-least:<pt>pt. Defaults to double."
        ),
    )
    return parser.parse_args()


def theme_is_safe(theme: object) -> bool:
    if theme is None:
        return True
    name = getattr(theme, "name", str(theme)).upper()
    return any(token in name for token in SAFE_THEME_TOKENS)


def inspect_font(font: Any) -> list[dict[str, str]]:
    findings: list[dict[str, str]] = []
    color = getattr(font, "color", None)
    if color is None:
        return findings

    rgb = getattr(color, "rgb", None)
    if rgb is not None and str(rgb).upper() not in {"000000", "AUTO"}:
        findings.append({"kind": "RGB", "value": str(rgb).upper()})

    theme = getattr(color, "theme_color", None)
    if theme is not None and not theme_is_safe(theme):
        findings.append({"kind": "THEME", "value": getattr(theme, "name", str(theme))})
    return findings


def has_border(element: Any) -> bool:
    if element is None:
        return False
    return bool(element.xpath(".//w:pBdr"))


def shading_fill(element: Any) -> str | None:
    if element is None:
        return None
    nodes = element.xpath(".//w:shd")
    for node in nodes:
        fill = node.get(qn("w:fill"))
        if fill and fill.upper() not in {"AUTO", "FFFFFF", "000000"}:
            return fill.upper()
    return None


def paragraph_is_heading(paragraph: Any, first_nonempty_index: int | None, index: int) -> bool:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if HEADING_STYLE_RE.match(style_name.strip()):
        return True
    return first_nonempty_index == index


def normalize_style_token(value: object) -> str:
    return re.sub(r"[\s_-]+", " ", str(value).strip().casefold())


def iter_style_chain(style: Any) -> Any:
    seen: set[str] = set()
    current = style
    while current is not None:
        style_id = str(getattr(current, "style_id", id(current)))
        if style_id in seen:
            break
        seen.add(style_id)
        yield current
        current = getattr(current, "base_style", None)


def paragraph_primary_style_tokens(paragraph: Any) -> set[str]:
    if paragraph.style is None:
        return set()
    return {
        normalize_style_token(paragraph.style.name),
        normalize_style_token(paragraph.style.style_id),
    }


def paragraph_is_list(paragraph: Any) -> bool:
    paragraph_properties = paragraph._p.pPr
    return paragraph_properties is not None and bool(
        paragraph_properties.xpath("./w:numPr")
    )


def paragraph_is_nonbody(paragraph: Any, excluded_styles: set[str]) -> bool:
    primary_tokens = paragraph_primary_style_tokens(paragraph)
    if primary_tokens & excluded_styles:
        return True
    if paragraph_is_list(paragraph):
        return True
    if paragraph.style is None:
        return False
    for style in iter_style_chain(paragraph.style):
        tokens = {
            normalize_style_token(style.name),
            normalize_style_token(style.style_id),
        }
        if any(NON_BODY_STYLE_RE.search(token) for token in tokens):
            return True
    return False


def paragraph_is_body(
    paragraph: Any,
    body_styles: set[str],
    excluded_styles: set[str],
) -> bool:
    if not paragraph.text.strip() or paragraph.style is None:
        return False

    if paragraph_is_nonbody(paragraph, excluded_styles):
        return False

    for style in iter_style_chain(paragraph.style):
        tokens = {
            normalize_style_token(style.name),
            normalize_style_token(style.style_id),
        }
        if tokens & body_styles:
            return True
    return False


def paragraph_is_structurally_empty(paragraph: Any) -> bool:
    if paragraph.text.strip():
        return False
    content_nodes = paragraph._p.xpath(
        ".//w:drawing | .//w:object | .//w:pict | .//w:br | "
        ".//w:fldChar | .//w:instrText | .//w:footnoteReference | "
        ".//w:endnoteReference"
    )
    return not content_nodes


def effective_spacing(paragraph: Any, field: str) -> tuple[float, str]:
    direct_value = getattr(paragraph.paragraph_format, field)
    if direct_value is not None:
        return float(direct_value.pt), "direct"

    if paragraph.style is not None:
        for style in iter_style_chain(paragraph.style):
            style_value = getattr(style.paragraph_format, field)
            if style_value is not None:
                return float(style_value.pt), f"style:{style.name}"
    return 0.0, "implicit-zero"


def automatic_spacing_sources(paragraph: Any, field: str) -> list[str]:
    attribute = {
        "space_before": "beforeAutospacing",
        "space_after": "afterAutospacing",
    }[field]
    sources: list[tuple[Any, str]] = [(paragraph._p.pPr, "direct")]
    if paragraph.style is not None:
        sources.extend(
            (style.element.pPr, f"style:{style.name}")
            for style in iter_style_chain(paragraph.style)
        )

    enabled_sources: list[str] = []
    for paragraph_properties, source in sources:
        if paragraph_properties is None:
            continue
        for spacing in paragraph_properties.xpath("./w:spacing"):
            value = spacing.get(qn(f"w:{attribute}"))
            if str(value or "").strip().casefold() in {"1", "true", "on"}:
                enabled_sources.append(source)
    return enabled_sources


def parse_line_spacing_spec(value: object) -> dict[str, object]:
    token = str(value).strip().casefold()
    if token in LINE_SPACING_ALIASES:
        kind, numeric = LINE_SPACING_ALIASES[token]
        return {"kind": kind, "value": numeric, "label": token}

    point_match = re.fullmatch(
        r"(exact|at-least)\s*:\s*([0-9]+(?:\.[0-9]+)?)\s*pt", token
    )
    if point_match:
        numeric = float(point_match.group(2))
        if numeric <= 0:
            raise ValueError("line spacing in points must be positive")
        return {
            "kind": point_match.group(1),
            "value": numeric,
            "label": f"{point_match.group(1)}:{numeric:g}pt",
        }

    try:
        numeric = float(token)
    except ValueError as exc:
        raise ValueError(
            "expected line spacing must be single, 1.15, 1.5, double, a positive "
            "numeric multiple, exact:<pt>pt, or at-least:<pt>pt"
        ) from exc
    if numeric <= 0:
        raise ValueError("line spacing multiple must be positive")
    return {"kind": "multiple", "value": numeric, "label": f"{numeric:g}"}


def effective_line_spacing(paragraph: Any) -> dict[str, object]:
    sources: list[tuple[Any, str]] = [
        (paragraph.paragraph_format, "direct")
    ]
    if paragraph.style is not None:
        sources.extend(
            (style.paragraph_format, f"style:{style.name}")
            for style in iter_style_chain(paragraph.style)
        )

    for paragraph_format, source in sources:
        value = paragraph_format.line_spacing
        if value is None:
            continue
        rule = paragraph_format.line_spacing_rule
        rule_name = getattr(rule, "name", str(rule)).casefold() if rule else ""
        if hasattr(value, "pt"):
            kind = "points"
            if "exact" in rule_name:
                kind = "exact"
            elif "at_least" in rule_name or "at least" in rule_name:
                kind = "at-least"
            return {
                "kind": kind,
                "value": float(value.pt),
                "source": source,
                "rule": rule_name or "points",
            }
        return {
            "kind": "multiple",
            "value": float(value),
            "source": source,
            "rule": rule_name or "multiple",
        }
    return {
        "kind": "implicit",
        "value": None,
        "source": "word-default",
        "rule": "implicit",
    }


def line_spacing_matches(
    actual: dict[str, object], expected: dict[str, object]
) -> bool:
    if actual["kind"] != expected["kind"]:
        return False
    actual_value = actual["value"]
    expected_value = expected["value"]
    if not isinstance(actual_value, (int, float)) or not isinstance(
        expected_value, (int, float)
    ):
        return False
    return abs(float(actual_value) - float(expected_value)) <= 0.01


def line_spacing_issue(
    paragraph: Any,
    index: int,
    role: str,
    code: str,
    expected: dict[str, object],
) -> dict[str, object] | None:
    actual = effective_line_spacing(paragraph)
    if line_spacing_matches(actual, expected):
        return None
    return {
        **paragraph_record(paragraph, index, role),
        "code": code,
        "detail": {
            "actual": actual,
            "expected": expected,
        },
    }


def top_level_blocks(document: Any) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    paragraph_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_index += 1
            blocks.append(
                {
                    "kind": "paragraph",
                    "paragraph": Paragraph(child, document),
                    "paragraph_index": paragraph_index,
                }
            )
        elif child.tag == qn("w:tbl"):
            blocks.append(
                {"kind": "table", "paragraph": None, "paragraph_index": None}
            )
    return blocks


def paragraph_record(paragraph: Any, index: int, role: str) -> dict[str, object]:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return {
        "paragraph": index,
        "role": role,
        "style": style_name,
        "text_preview": paragraph.text.strip().replace("\n", " ")[:100],
    }


def field_instruction(node: Any) -> str:
    if node.tag == qn("w:fldSimple"):
        return str(node.get(qn("w:instr")) or "")
    return str(node.text or "")


def story_has_page_field(story: Any) -> bool:
    for node in story._element.xpath(".//w:instrText | .//w:fldSimple"):
        if PAGE_FIELD_RE.search(field_instruction(node)):
            return True
    return False


def audit_document_numbering(
    document: Any,
) -> tuple[list[dict[str, object]], dict[str, object]]:
    issues: list[dict[str, object]] = []
    sections = list(document.sections)
    line_numbered_sections = 0

    for section_index, section in enumerate(sections, start=1):
        line_number_nodes = section._sectPr.xpath("./w:lnNumType")
        if not line_number_nodes:
            issues.append(
                {
                    "section": section_index,
                    "role": "section-numbering",
                    "code": "SECTION_LINE_NUMBERING_MISSING",
                    "detail": "Every section requires Word-native line numbering.",
                }
            )
            continue

        if len(line_number_nodes) > 1:
            issues.append(
                {
                    "section": section_index,
                    "role": "section-numbering",
                    "code": "DUPLICATE_SECTION_LINE_NUMBERING",
                    "detail": {"found": len(line_number_nodes), "required": 1},
                }
            )

        line_numbers = line_number_nodes[0]
        count_by = str(line_numbers.get(qn("w:countBy")) or "").strip()
        restart = str(line_numbers.get(qn("w:restart")) or "").strip().casefold()
        start = str(line_numbers.get(qn("w:start")) or "").strip()
        section_pass = True
        if count_by != "1":
            section_pass = False
            issues.append(
                {
                    "section": section_index,
                    "role": "section-numbering",
                    "code": "SECTION_LINE_NUMBERING_COUNTBY_NOT_ONE",
                    "detail": {"actual": count_by or "implicit", "required": "1"},
                }
            )
        if restart != "continuous":
            section_pass = False
            issues.append(
                {
                    "section": section_index,
                    "role": "section-numbering",
                    "code": "SECTION_LINE_NUMBERING_NOT_CONTINUOUS",
                    "detail": {
                        "actual": restart or "implicit",
                        "required": "continuous",
                    },
                }
            )
        if start and start != "1":
            section_pass = False
            issues.append(
                {
                    "section": section_index,
                    "role": "section-numbering",
                    "code": "SECTION_LINE_NUMBERING_START_INVALID",
                    "detail": {"actual": start, "allowed": ["implicit", "1"]},
                }
            )
        if section_pass:
            line_numbered_sections += 1

    suppressed_nodes = document.element.body.xpath(
        ".//w:pPr/w:suppressLineNumbers"
    )
    if suppressed_nodes:
        issues.append(
            {
                "role": "document-numbering",
                "code": "LINE_NUMBER_SUPPRESSION_PRESENT",
                "detail": {
                    "suppressed_paragraph_count": len(suppressed_nodes),
                    "required": 0,
                },
            }
        )

    even_and_odd = (
        document.settings.element.find(qn("w:evenAndOddHeaders")) is not None
    )
    required_page_stories = 0
    page_stories_with_field = 0
    page_restart_sections: list[int] = []

    for section_index, section in enumerate(sections, start=1):
        stories = [
            ("default", section.header, section.footer),
        ]
        if section.different_first_page_header_footer:
            stories.append(
                (
                    "first",
                    section.first_page_header,
                    section.first_page_footer,
                )
            )
        if even_and_odd:
            stories.append(("even", section.even_page_header, section.even_page_footer))

        for story_name, header, footer in stories:
            required_page_stories += 1
            if story_has_page_field(header) or story_has_page_field(footer):
                page_stories_with_field += 1
            else:
                issues.append(
                    {
                        "section": section_index,
                        "role": "page-numbering",
                        "code": "PAGE_NUMBER_FIELD_MISSING",
                        "detail": {
                            "story": story_name,
                            "required": "dynamic PAGE field in active header or footer",
                        },
                    }
                )

        for page_number_node in section._sectPr.xpath("./w:pgNumType"):
            start = str(page_number_node.get(qn("w:start")) or "").strip()
            if start and (section_index > 1 or start != "1"):
                page_restart_sections.append(section_index)
                issues.append(
                    {
                        "section": section_index,
                        "role": "page-numbering",
                        "code": "PAGE_NUMBERING_NOT_CONTINUOUS",
                        "detail": {
                            "start": start,
                            "required": "continuous across sections",
                        },
                    }
                )

    summary = {
        "section_count": len(sections),
        "line_numbering": "continuous",
        "line_numbered_section_count": line_numbered_sections,
        "suppressed_paragraph_count": len(suppressed_nodes),
        "page_numbering": "continuous",
        "required_page_story_count": required_page_stories,
        "page_story_with_field_count": page_stories_with_field,
        "page_restart_sections": page_restart_sections,
        "even_and_odd_headers_enabled": even_and_odd,
    }
    return issues, summary


def audit_body_paragraph_rhythm(
    document: Any,
    body_styles: set[str],
    excluded_styles: set[str],
    paragraph_separation: str,
    expected_line_spacing: dict[str, object],
) -> tuple[list[dict[str, object]], dict[str, object]]:
    issues: list[dict[str, object]] = []
    blocks = top_level_blocks(document)
    body_flags = [False] * len(blocks)
    inspected_body_count = 0
    body_pair_count = 0
    excluded_nonbody_count = 0
    unclassified_count = 0
    body_style_names: set[str] = set()
    excluded_style_names: set[str] = set()

    for block_index, block in enumerate(blocks):
        if block["kind"] != "paragraph":
            continue
        paragraph = block["paragraph"]
        assert paragraph is not None
        if not paragraph.text.strip():
            continue

        if not paragraph_is_body(paragraph, body_styles, excluded_styles):
            if paragraph_is_nonbody(paragraph, excluded_styles):
                excluded_nonbody_count += 1
                style_name = paragraph.style.name if paragraph.style is not None else ""
                excluded_style_names.add(style_name)
                continue
            unclassified_count += 1
            issues.append(
                {
                    **paragraph_record(
                        paragraph,
                        int(block["paragraph_index"]),
                        "unclassified-nonempty",
                    ),
                    "code": "UNCLASSIFIED_NONEMPTY_PARAGRAPH_STYLE",
                    "detail": (
                        "Every non-empty top-level paragraph must be classified as "
                        "body prose with --body-style or as non-body with "
                        "--exclude-style before DOCX delivery."
                    ),
                }
            )
            continue

        body_flags[block_index] = True
        inspected_body_count += 1
        style_name = paragraph.style.name if paragraph.style is not None else ""
        body_style_names.add(style_name)
        record = paragraph_record(
            paragraph, int(block["paragraph_index"]), "body-prose"
        )
        if paragraph_separation == "literal-blank":
            for field, label in (
                ("space_before", "BEFORE"),
                ("space_after", "AFTER"),
            ):
                automatic_sources = automatic_spacing_sources(paragraph, field)
                if automatic_sources:
                    issues.append(
                        {
                            **record,
                            "code": f"BODY_PARAGRAPH_AUTOSPACING_{label}",
                            "detail": {
                                "enabled_sources": automatic_sources,
                                "required": "disabled",
                            },
                        }
                    )
                points, source = effective_spacing(paragraph, field)
                if abs(points) > 0.01:
                    issues.append(
                        {
                            **record,
                            "code": f"BODY_PARAGRAPH_SPACE_{label}",
                            "detail": {
                                "points": round(points, 3),
                                "source": source,
                                "required_points": 0,
                            },
                        }
                    )
        line_issue = line_spacing_issue(
            paragraph,
            int(block["paragraph_index"]),
            "body-prose",
            "BODY_LINE_SPACING_MISMATCH",
            expected_line_spacing,
        )
        if line_issue:
            issues.append(line_issue)

    for block_index, is_body in enumerate(body_flags):
        if not is_body:
            continue
        separator_blocks: list[dict[str, object]] = []
        next_index = block_index + 1
        while next_index < len(blocks):
            next_block = blocks[next_index]
            if next_block["kind"] != "paragraph":
                break
            next_paragraph = next_block["paragraph"]
            assert next_paragraph is not None
            if not paragraph_is_structurally_empty(next_paragraph):
                break
            separator_blocks.append(next_block)
            next_index += 1

        if next_index >= len(blocks) or not body_flags[next_index]:
            continue

        body_pair_count += 1
        paragraph = blocks[block_index]["paragraph"]
        assert paragraph is not None
        record = paragraph_record(
            paragraph,
            int(blocks[block_index]["paragraph_index"]),
            "body-prose",
        )

        if not separator_blocks:
            issues.append(
                {
                    **record,
                    "code": "MISSING_LITERAL_BLANK_PARAGRAPH",
                    "detail": (
                        "Adjacent body-prose paragraphs require exactly one real "
                        "empty paragraph, not paragraph spacing."
                    ),
                }
            )
            continue

        if len(separator_blocks) > 1:
            issues.append(
                {
                    **record,
                    "code": "MULTIPLE_LITERAL_BLANK_PARAGRAPHS",
                    "detail": {
                        "found": len(separator_blocks),
                        "required": 1,
                    },
                }
            )

        for separator in separator_blocks:
            blank = separator["paragraph"]
            assert blank is not None
            blank_record = paragraph_record(
                blank, int(separator["paragraph_index"]), "blank-separator"
            )
            for field, label in (
                ("space_before", "BEFORE"),
                ("space_after", "AFTER"),
            ):
                automatic_sources = automatic_spacing_sources(blank, field)
                if automatic_sources:
                    issues.append(
                        {
                            **blank_record,
                            "code": f"BLANK_PARAGRAPH_AUTOSPACING_{label}",
                            "detail": {
                                "enabled_sources": automatic_sources,
                                "required": "disabled",
                            },
                        }
                    )
                points, source = effective_spacing(blank, field)
                if abs(points) > 0.01:
                    issues.append(
                        {
                            **blank_record,
                            "code": f"BLANK_PARAGRAPH_SPACE_{label}",
                            "detail": {
                                "points": round(points, 3),
                                "source": source,
                                "required_points": 0,
                            },
                        }
                    )
            line_issue = line_spacing_issue(
                blank,
                int(separator["paragraph_index"]),
                "blank-separator",
                "BLANK_PARAGRAPH_LINE_SPACING_MISMATCH",
                expected_line_spacing,
            )
            if line_issue:
                issues.append(line_issue)

    summary = {
        "mode": paragraph_separation,
        "expected_line_spacing": expected_line_spacing["label"],
        "inspected_body_paragraph_count": inspected_body_count,
        "adjacent_body_pair_count": body_pair_count,
        "excluded_nonbody_paragraph_count": excluded_nonbody_count,
        "unclassified_nonempty_paragraph_count": unclassified_count,
        "body_style_names": sorted(body_style_names),
        "excluded_nonbody_style_names": sorted(excluded_style_names),
    }
    return issues, summary


def audit(
    path: Path,
    paragraph_separation: str = "literal-blank",
    extra_body_styles: list[str] | None = None,
    excluded_styles: list[str] | None = None,
    expected_line_spacing: object = "double",
) -> dict[str, object]:
    if paragraph_separation != "literal-blank":
        raise ValueError(
            "paragraph_separation must be literal-blank; this Skill does not "
            "permit a DOCX-output bypass"
        )
    document = Document(str(path))
    issues: list[dict[str, object]] = []
    inspected: list[dict[str, object]] = []

    first_nonempty_index: int | None = None
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.text.strip():
            first_nonempty_index = index
            break

    for index, paragraph in enumerate(document.paragraphs, start=1):
        if not paragraph.text.strip():
            continue
        if not paragraph_is_heading(paragraph, first_nonempty_index, index):
            continue

        style_name = paragraph.style.name if paragraph.style is not None else ""
        role = (
            "first-paragraph-title-candidate"
            if index == first_nonempty_index and not HEADING_STYLE_RE.match(style_name.strip())
            else style_name
        )
        preview = paragraph.text.strip().replace("\n", " ")[:100]
        record = {
            "paragraph": index,
            "role": role,
            "style": style_name,
            "text_preview": preview,
        }
        inspected.append(record)

        if paragraph.style is not None:
            for color in inspect_font(paragraph.style.font):
                issues.append(
                    {
                        **record,
                        "code": "NON_BLACK_STYLE_COLOR",
                        "detail": color,
                    }
                )

        for run_index, run in enumerate(paragraph.runs, start=1):
            for color in inspect_font(run.font):
                issues.append(
                    {
                        **record,
                        "run": run_index,
                        "code": "NON_BLACK_DIRECT_COLOR",
                        "detail": color,
                    }
                )

        paragraph_ppr = paragraph._p.pPr
        style_ppr = paragraph.style.element.pPr if paragraph.style is not None else None
        if has_border(paragraph_ppr) or has_border(style_ppr):
            issues.append(
                {
                    **record,
                    "code": "DECORATIVE_PARAGRAPH_BORDER",
                    "detail": "Title/heading paragraph has a border or rule.",
                }
            )

        fill = shading_fill(paragraph_ppr) or shading_fill(style_ppr)
        if fill:
            issues.append(
                {
                    **record,
                    "code": "TITLE_OR_HEADING_SHADING",
                    "detail": f"Non-neutral shading fill {fill}.",
                }
            )

    body_styles = set(DEFAULT_BODY_STYLES)
    body_styles.update(
        normalize_style_token(value) for value in (extra_body_styles or [])
    )
    excluded_style_tokens = {
        normalize_style_token(value) for value in (excluded_styles or [])
    }
    line_spacing_spec = parse_line_spacing_spec(expected_line_spacing)
    paragraph_structure: dict[str, object] = {
        "mode": paragraph_separation,
        "expected_line_spacing": line_spacing_spec["label"],
        "inspected_body_paragraph_count": 0,
        "adjacent_body_pair_count": 0,
        "excluded_nonbody_paragraph_count": 0,
        "unclassified_nonempty_paragraph_count": 0,
        "body_style_names": [],
        "excluded_nonbody_style_names": [],
    }
    structure_issues, paragraph_structure = audit_body_paragraph_rhythm(
        document,
        body_styles,
        excluded_style_tokens,
        paragraph_separation,
        line_spacing_spec,
    )
    issues.extend(structure_issues)
    numbering_issues, document_numbering = audit_document_numbering(document)
    issues.extend(numbering_issues)

    return {
        "status": "FAIL" if issues else "MECHANICAL_PASS",
        "document": str(path.resolve()),
        "paragraph_count": len(document.paragraphs),
        "inspected_title_heading_count": len(inspected),
        "paragraph_structure": paragraph_structure,
        "document_numbering": document_numbering,
        "issue_count": len(issues),
        "issues": issues,
        "inspected": inspected,
        "boundary": (
            "Mechanical style, paragraph-structure, and Word numbering screening only. "
            "Official-template review and rendered page-by-page visual QA are still "
            "required."
        ),
    }


def render_text(report: dict[str, object]) -> str:
    lines = [
        f"DOCX manuscript style audit: {report['status']}",
        f"Document: {report['document']}",
        f"Inspected title/headings: {report['inspected_title_heading_count']}",
        "Paragraph separation: "
        f"{report['paragraph_structure']['mode']} "
        f"({report['paragraph_structure']['inspected_body_paragraph_count']} body "
        "paragraphs)",
        "Expected body line spacing: "
        f"{report['paragraph_structure']['expected_line_spacing']}",
        "Unclassified non-empty paragraphs: "
        f"{report['paragraph_structure']['unclassified_nonempty_paragraph_count']}",
        "Continuous line-numbered sections: "
        f"{report['document_numbering']['line_numbered_section_count']}/"
        f"{report['document_numbering']['section_count']}",
        "Dynamic PAGE fields in active page stories: "
        f"{report['document_numbering']['page_story_with_field_count']}/"
        f"{report['document_numbering']['required_page_story_count']}",
        f"Issues: {report['issue_count']}",
    ]
    for issue in report["issues"]:
        if "paragraph" in issue:
            location = f"paragraph {issue['paragraph']}"
        elif "section" in issue:
            location = f"section {issue['section']}"
        else:
            location = "document"
        role = issue.get("role", "document")
        lines.append(
            f"ERROR: {location} ({role}): "
            f"{issue['code']} - {issue['detail']}"
        )
    lines.append(f"BOUNDARY: {report['boundary']}")
    return "\n".join(lines)


def main() -> int:
    args = parse_args()
    try:
        report = audit(
            args.document,
            paragraph_separation=args.paragraph_separation,
            extra_body_styles=args.body_style,
            excluded_styles=args.exclude_style,
            expected_line_spacing=args.expected_line_spacing,
        )
    except (OSError, ValueError) as exc:
        print(f"Unable to audit DOCX: {exc}", file=sys.stderr)
        return 2

    if args.output_json:
        args.output_json.write_text(
            json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
        )
    if args.json:
        print(json.dumps(report, indent=2, ensure_ascii=False))
    else:
        print(render_text(report))
    return 0 if report["status"] == "MECHANICAL_PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
